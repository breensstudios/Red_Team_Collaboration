from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pymysql
import json
import os
from datetime import datetime, timedelta
import hashlib
import re
import uuid
import socket
import requests
from bs4 import BeautifulSoup
import urllib3
import jwt
from functools import wraps

# 禁用SSL警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

app = Flask(__name__)
CORS(app)

CONFIG_FILE = os.environ.get('CONFIG_FILE', 'config.json')
# JWT密钥 - 生产环境应该从环境变量或配置文件读取
JWT_SECRET_KEY = os.environ.get('JWT_SECRET_KEY', 'your-secret-key-change-in-production-' + str(uuid.uuid4()))
JWT_ALGORITHM = 'HS256'
JWT_EXPIRATION_HOURS = 24

def load_config():
    """加载配置文件"""
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None

def save_config(config):
    """保存配置文件"""
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

def get_db_connection():
    """获取数据库连接"""
    config = load_config()
    if not config:
        return None
    
    try:
        connection = pymysql.connect(
            host=config['host'],
            port=int(config['port']),
            user=config['user'],
            password=config['password'],
            database=config['database'],
            charset='utf8mb4',
            cursorclass=pymysql.cursors.DictCursor
        )
        return connection
    except Exception as e:
        print(f"Database connection error: {e}")
        return None

def check_installed():
    """检查是否已安装"""
    config = load_config()
    if not config:
        return False
    
    try:
        conn = get_db_connection()
        if not conn:
            return False
        
        cursor = conn.cursor()
        cursor.execute("SHOW TABLES LIKE 'users'")
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        return result is not None
    except:
        return False

def generate_token(user_id, username, is_admin=False, is_super_admin=False):
    """生成JWT token"""
    payload = {
        'user_id': user_id,
        'username': username,
        'is_admin': is_admin,
        'is_super_admin': is_super_admin,
        'exp': datetime.utcnow() + timedelta(hours=JWT_EXPIRATION_HOURS),
        'iat': datetime.utcnow()
    }
    token = jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
    return token

def verify_token(token):
    """验证JWT token并返回用户信息"""
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        return {
            'id': payload['user_id'],
            'username': payload['username'],
            'is_admin': payload.get('is_admin', False),
            'is_super_admin': payload.get('is_super_admin', False)
        }
    except jwt.ExpiredSignatureError:
        return None  # Token已过期
    except jwt.InvalidTokenError:
        return None  # Token无效

def require_auth(f):
    """装饰器：要求用户认证"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 从Authorization header获取token
        auth_header = request.headers.get('Authorization')
        if not auth_header:
            return jsonify({'success': False, 'message': '未提供认证令牌'}), 401
        
        # 格式: "Bearer <token>"
        parts = auth_header.split()
        if len(parts) != 2 or parts[0].lower() != 'bearer':
            return jsonify({'success': False, 'message': '认证令牌格式错误'}), 401
        
        token = parts[1]
        user = verify_token(token)
        
        if not user:
            return jsonify({'success': False, 'message': '认证令牌无效或已过期'}), 401
        
        # 将用户信息附加到request对象
        request.current_user = user
        return f(*args, **kwargs)
    
    return decorated_function

def require_admin(f):
    """装饰器：要求管理员权限"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
        
        if not user.get('is_admin') and not user.get('is_super_admin'):
            return jsonify({'success': False, 'message': '需要管理员权限'}), 403
        
        return f(*args, **kwargs)
    
    return decorated_function

def require_super_admin(f):
    """装饰器：要求超级管理员权限"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        user = get_current_user()
        if not user:
            return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
        
        if not user.get('is_super_admin'):
            return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
        
        return f(*args, **kwargs)
    
    return decorated_function

def get_current_user():
    """获取当前登录用户（仅支持JWT认证）"""
    # 从Authorization header获取JWT token
    auth_header = request.headers.get('Authorization')
    if not auth_header:
        return None
    
    parts = auth_header.split()
    if len(parts) != 2 or parts[0].lower() != 'bearer':
        return None
    
    token = parts[1]
    user = verify_token(token)
    if not user:
        return None
    
    # 为了确保权限信息最新，从数据库重新查询用户信息
    conn = get_db_connection()
    if not conn:
        return None
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, username, is_admin, is_super_admin FROM users WHERE id = %s",
            (user['id'],)
        )
        db_user = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if db_user:
            return {
                'id': db_user['id'],
                'username': db_user['username'],
                'is_admin': db_user.get('is_admin', False),
                'is_super_admin': db_user.get('is_super_admin', False)
            }
    except:
        pass
    
    return None

def check_project_permission(cursor, project_id, user_id, require_owner=False):
    """
    检查用户对项目的访问权限
    :param cursor: 数据库游标
    :param project_id: 项目ID
    :param user_id: 用户ID
    :param require_owner: 是否要求必须是项目所有者
    :return: True表示有权限，False表示无权限
    """
    cursor.execute("""
        SELECT created_by FROM projects WHERE id = %s
    """, (project_id,))
    project = cursor.fetchone()
    
    if not project:
        return False
    
    # 如果要求必须是所有者
    if require_owner:
        return project['created_by'] == user_id
    
    # 检查是否是项目创建者
    if project['created_by'] == user_id:
        return True
    
    # 检查是否是协作者且项目可见
    cursor.execute("""
        SELECT 1 FROM user_collaborations uc
        JOIN projects p ON p.created_by = uc.collaborator_id
        WHERE uc.user_id = %s 
          AND p.id = %s
          AND (p.is_visible_to_inviter IS NULL OR p.is_visible_to_inviter = TRUE)
    """, (user_id, project_id))
    
    return cursor.fetchone() is not None

def check_asset_permission(cursor, asset_id, user_id, require_owner=False):
    """
    检查用户对资产的访问权限
    :param cursor: 数据库游标
    :param asset_id: 资产ID
    :param user_id: 用户ID
    :param require_owner: 是否要求必须是资产所有者
    :return: True表示有权限，False表示无权限
    """
    cursor.execute("""
        SELECT a.created_by, a.project_id, p.created_by as project_owner
        FROM assets a
        JOIN projects p ON a.project_id = p.id
        WHERE a.id = %s
    """, (asset_id,))
    asset = cursor.fetchone()
    
    if not asset:
        return False
    
    # 如果要求必须是所有者
    if require_owner:
        return asset['created_by'] == user_id
    
    # 检查是否是资产创建者
    if asset['created_by'] == user_id:
        return True
    
    # 检查是否是项目所有者
    if asset['project_owner'] == user_id:
        return True
    
    # 检查项目权限
    return check_project_permission(cursor, asset['project_id'], user_id, False)

def check_vulnerability_permission(cursor, vuln_id, user_id, require_owner=False):
    """
    检查用户对漏洞的访问权限
    """
    cursor.execute("""
        SELECT v.created_by, v.project_id, p.created_by as project_owner
        FROM vulnerabilities v
        JOIN projects p ON v.project_id = p.id
        WHERE v.id = %s
    """, (vuln_id,))
    vuln = cursor.fetchone()
    
    if not vuln:
        return False
    
    if require_owner:
        return vuln['created_by'] == user_id
    
    if vuln['created_by'] == user_id:
        return True
    
    if vuln['project_owner'] == user_id:
        return True
    
    return check_project_permission(cursor, vuln['project_id'], user_id, False)

def check_curl_task_permission(cursor, task_id, user_id, require_owner=False):
    """
    检查用户对curl任务的访问权限
    """
    cursor.execute("""
        SELECT ct.created_by, ct.project_id, p.created_by as project_owner
        FROM curl_tasks ct
        JOIN projects p ON ct.project_id = p.id
        WHERE ct.id = %s
    """, (task_id,))
    task = cursor.fetchone()
    
    if not task:
        return False
    
    if require_owner:
        return task['created_by'] == user_id
    
    if task['created_by'] == user_id:
        return True
    
    if task['project_owner'] == user_id:
        return True
    
    return check_project_permission(cursor, task['project_id'], user_id, False)

@app.route('/api/check-install', methods=['GET'])
def check_install():
    """检查安装状态"""
    installed = check_installed()
    return jsonify({'installed': installed})

@app.route('/api/install', methods=['POST'])
def install():
    """安装系统，创建数据库表结构"""
    if check_installed():
        return jsonify({'success': False, 'message': '系统已安装，无法重复安装'}), 400
    
    data = request.json
    host = data.get('host')
    port = data.get('port', 3306)
    user = data.get('user')
    password = data.get('password')
    database = data.get('database')
    
    if not all([host, user, password, database]):
        return jsonify({'success': False, 'message': '请填写完整的数据库信息'}), 400
    
    # 验证数据库名称，防止SQL注入
    if not re.match(r'^[a-zA-Z0-9_]+$', database):
        return jsonify({'success': False, 'message': '数据库名称只能包含字母、数字和下划线'}), 400
    
    try:
        # 先连接MySQL服务器（不指定数据库）
        connection = pymysql.connect(
            host=host,
            port=int(port),
            user=user,
            password=password,
            charset='utf8mb4',
            cursorclass=pymysql.cursors.DictCursor
        )
        
        cursor = connection.cursor()
        
        # 创建数据库（如果不存在）- 使用反引号包裹，并已验证输入
        cursor.execute(f"CREATE DATABASE IF NOT EXISTS `{database}` CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci")
        cursor.execute(f"USE `{database}`")
        
        # 创建用户表（包含管理员字段）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `users` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `username` VARCHAR(50) NOT NULL UNIQUE,
                `password` VARCHAR(64) NOT NULL,
                `is_admin` BOOLEAN DEFAULT FALSE COMMENT '是否为管理员',
                `is_super_admin` BOOLEAN DEFAULT FALSE COMMENT '是否为超级管理员',
                `fofa_email` VARCHAR(200) DEFAULT NULL COMMENT 'FOFA账号Email',
                `fofa_api_key` VARCHAR(200) DEFAULT NULL COMMENT 'FOFA API Key',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                INDEX idx_username (username)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        
        # 创建项目表（包含协同任务字段）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `projects` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `name` VARCHAR(100) NOT NULL,
                `description` TEXT,
                `created_by` INT NOT NULL,
                `is_collaboration_task` BOOLEAN DEFAULT FALSE COMMENT '是否为协同任务项目',
                `source_project_id` INT DEFAULT NULL COMMENT '源项目ID（协同任务时使用）',
                `task_assigner_id` INT DEFAULT NULL COMMENT '任务下发人ID',
                `is_visible_to_inviter` BOOLEAN DEFAULT TRUE COMMENT '是否对邀请人可见',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_created_by (created_by),
                INDEX idx_is_collaboration_task (is_collaboration_task),
                INDEX idx_source_project_id (source_project_id),
                INDEX idx_task_assigner_id (task_assigner_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        
        # 创建资产表（包含源资产ID字段）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `assets` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `project_id` INT NOT NULL,
                `asset_type` ENUM('ip', 'ip_port', 'http', 'https', 'domain') NOT NULL,
                `asset_value` VARCHAR(255) NOT NULL,
                `ip` VARCHAR(45),
                `port` INT,
                `protocol` VARCHAR(10),
                `url` TEXT,
                `title` VARCHAR(500) DEFAULT NULL COMMENT '网页标题',
                `content_length` INT DEFAULT NULL COMMENT '网页内容长度',
                `status` ENUM('未测试', '信息收集', '探测中', '存在漏洞', '已验证可打点') DEFAULT '未测试',
                `risk_level` ENUM('低风险', '高危', '重点打点', 'CDN', 'WAF') DEFAULT NULL,
                `tags` TEXT,
                `notes` TEXT,
                `test_result` TEXT COMMENT '测试结果（Markdown格式）',
                `source_asset_id` INT DEFAULT NULL COMMENT '源资产ID（协同任务同步时使用）',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_project_id (project_id),
                INDEX idx_asset_type (asset_type),
                INDEX idx_ip (ip),
                INDEX idx_status (status),
                INDEX idx_risk_level (risk_level),
                INDEX idx_project_type (project_id, asset_type),
                INDEX idx_project_status (project_id, status),
                INDEX idx_project_risk (project_id, risk_level),
                INDEX idx_project_created (project_id, created_at DESC),
                INDEX idx_asset_value (asset_value(100)),
                INDEX idx_port (port),
                INDEX idx_protocol (protocol),
                INDEX idx_created_by (created_by),
                INDEX idx_project_ip (project_id, ip),
                INDEX idx_title (title(100)),
                INDEX idx_source_asset (source_asset_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        
        # 创建邀请码表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `invite_codes` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `code` VARCHAR(32) NOT NULL UNIQUE,
                `user_id` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `expires_at` DATETIME,
                `is_used` BOOLEAN DEFAULT FALSE,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_code (code),
                INDEX idx_user_id (user_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        
        # 创建用户协作关系表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `user_collaborations` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `user_id` INT NOT NULL,
                `collaborator_id` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
                FOREIGN KEY (collaborator_id) REFERENCES users(id) ON DELETE CASCADE,
                UNIQUE KEY unique_collaboration (user_id, collaborator_id),
                INDEX idx_user_id (user_id),
                INDEX idx_collaborator_id (collaborator_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        
        # 创建共享凭证表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `shared_credentials` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `platform` VARCHAR(100) NOT NULL COMMENT '平台名称',
                `url` VARCHAR(500) DEFAULT NULL COMMENT '平台URL',
                `username` VARCHAR(200) NOT NULL COMMENT '用户名/账号',
                `password` TEXT NOT NULL COMMENT '密码（加密存储）',
                `category` ENUM('VPN', '云服务', '开发工具', '测试平台', '办公软件', '其他') DEFAULT '其他' COMMENT '分类',
                `description` TEXT DEFAULT NULL COMMENT '备注说明',
                `is_active` BOOLEAN DEFAULT TRUE COMMENT '是否有效',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_platform (platform),
                INDEX idx_category (category),
                INDEX idx_created_by (created_by),
                INDEX idx_is_active (is_active)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='共享凭证表'
        """)
        
        # 创建curl任务表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `curl_tasks` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `project_id` INT NOT NULL,
                `name` VARCHAR(200) NOT NULL COMMENT '任务名称',
                `template_type` ENUM('custom', 'lighthouse', 'fofa') DEFAULT 'custom' COMMENT '模板类型',
                `template_config` TEXT DEFAULT NULL COMMENT '模板配置（JSON）',
                `curl_command` TEXT DEFAULT NULL COMMENT 'curl命令',
                `extract_pattern` VARCHAR(500) DEFAULT NULL COMMENT '提取模式（正则表达式）',
                `schedule_type` ENUM('manual', 'interval', 'cron') DEFAULT 'manual' COMMENT '调度类型',
                `schedule_config` TEXT DEFAULT NULL COMMENT '调度配置（JSON）',
                `enabled` BOOLEAN DEFAULT TRUE COMMENT '是否启用',
                `last_run_time` DATETIME DEFAULT NULL COMMENT '上次运行时间',
                `last_run_status` ENUM('success', 'failed', 'running') DEFAULT NULL COMMENT '上次运行状态',
                `last_run_result` MEDIUMTEXT DEFAULT NULL COMMENT '上次运行结果',
                `assets_extracted` INT DEFAULT 0 COMMENT '提取的资产数量',
                `batch_tags` TEXT DEFAULT NULL COMMENT '批量标签（JSON数组）',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_project_id (project_id),
                INDEX idx_enabled (enabled),
                INDEX idx_schedule_type (schedule_type),
                INDEX idx_template_type (template_type)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='curl任务表'
        """)
        
        # 创建curl任务执行日志表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `curl_task_logs` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `task_id` INT NOT NULL,
                `status` ENUM('success', 'failed') NOT NULL COMMENT '执行状态',
                `response_content` LONGTEXT DEFAULT NULL COMMENT '响应内容',
                `assets_extracted` INT DEFAULT 0 COMMENT '提取的资产数量',
                `error_message` TEXT DEFAULT NULL COMMENT '错误信息',
                `execution_time` INT DEFAULT 0 COMMENT '执行时间（毫秒）',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (task_id) REFERENCES curl_tasks(id) ON DELETE CASCADE,
                INDEX idx_task_id (task_id),
                INDEX idx_status (status),
                INDEX idx_created_at (created_at)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='curl任务执行日志表'
        """)
        
        # 创建漏洞表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `vulnerabilities` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `project_id` INT NOT NULL,
                `url` VARCHAR(500) NOT NULL COMMENT '漏洞URL',
                `vuln_type` VARCHAR(100) NOT NULL COMMENT '漏洞类型',
                `description` TEXT DEFAULT NULL COMMENT '漏洞描述',
                `discoverer` VARCHAR(100) DEFAULT NULL COMMENT '发现人员',
                `markdown_detail` MEDIUMTEXT DEFAULT NULL COMMENT 'Markdown详情',
                `status` ENUM('未提交', '已提交', '未修复', '已修复', '无法复现', '无法访问') DEFAULT '未提交' COMMENT '漏洞状态',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_project_id (project_id),
                INDEX idx_vuln_type (vuln_type),
                INDEX idx_status (status),
                INDEX idx_discoverer (discoverer),
                INDEX idx_created_by (created_by),
                INDEX idx_project_status (project_id, status)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='漏洞生命周期管理表'
        """)
        
        # 创建系统设置表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `system_settings` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `setting_key` VARCHAR(100) NOT NULL UNIQUE COMMENT '设置键',
                `setting_value` TEXT NOT NULL COMMENT '设置值',
                `description` VARCHAR(500) DEFAULT NULL COMMENT '设置描述',
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                INDEX idx_setting_key (setting_key)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='系统设置表'
        """)
        
        # 插入默认系统设置
        default_settings = [
            ('allow_registration', 'true', '是否允许用户注册'),
            ('site_name', '红队资产协作平台', '网站名称'),
            ('site_description', 'RED team Asset Collaboration Platform', '网站描述')
        ]
        
        for key, value, desc in default_settings:
            cursor.execute("""
                INSERT INTO system_settings (setting_key, setting_value, description)
                VALUES (%s, %s, %s)
                ON DUPLICATE KEY UPDATE 
                    setting_value = VALUES(setting_value),
                    description = VALUES(description)
            """, (key, value, desc))
        
        # 创建报告模板表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `report_templates` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `name` VARCHAR(200) NOT NULL COMMENT '模板名称',
                `description` TEXT DEFAULT NULL COMMENT '模板描述',
                `file_name` VARCHAR(255) NOT NULL COMMENT '文件名',
                `file_path` VARCHAR(500) NOT NULL COMMENT '文件存储路径',
                `file_size` INT NOT NULL COMMENT '文件大小（字节）',
                `placeholders` TEXT DEFAULT NULL COMMENT '占位符列表（JSON）',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_created_by (created_by)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='报告模板表'
        """)
        
        # 创建报告实例表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `reports` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `template_id` INT NOT NULL COMMENT '模板ID',
                `project_id` INT DEFAULT NULL COMMENT '关联项目ID',
                `name` VARCHAR(200) NOT NULL COMMENT '报告名称',
                `content` MEDIUMTEXT DEFAULT NULL COMMENT '报告内容（JSON）',
                `status` ENUM('draft', 'completed') DEFAULT 'draft' COMMENT '报告状态',
                `created_by` INT NOT NULL,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (template_id) REFERENCES report_templates(id) ON DELETE CASCADE,
                FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE SET NULL,
                FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_template_id (template_id),
                INDEX idx_project_id (project_id),
                INDEX idx_created_by (created_by),
                INDEX idx_status (status)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='报告实例表'
        """)
        
        # 创建报告附件表（存储图片等）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `report_attachments` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `report_id` INT NOT NULL COMMENT '报告ID',
                `file_name` VARCHAR(255) NOT NULL COMMENT '文件名',
                `file_path` VARCHAR(500) NOT NULL COMMENT '文件路径',
                `file_type` VARCHAR(50) NOT NULL COMMENT '文件类型',
                `file_size` INT NOT NULL COMMENT '文件大小',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE,
                INDEX idx_report_id (report_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='报告附件表'
        """)
        
        # 创建协同任务分配表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `collaboration_tasks` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `source_project_id` INT NOT NULL COMMENT '源项目ID',
                `target_project_id` INT NOT NULL COMMENT '目标项目ID（被分配人的项目）',
                `assigner_id` INT NOT NULL COMMENT '任务下发人ID',
                `assignee_id` INT NOT NULL COMMENT '任务接收人ID',
                `asset_ids` TEXT DEFAULT NULL COMMENT '下发的资产ID列表（JSON数组）',
                `task_description` TEXT DEFAULT NULL COMMENT '任务描述',
                `status` ENUM('active', 'completed', 'cancelled') DEFAULT 'active' COMMENT '任务状态',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (source_project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (target_project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (assigner_id) REFERENCES users(id) ON DELETE CASCADE,
                FOREIGN KEY (assignee_id) REFERENCES users(id) ON DELETE CASCADE,
                INDEX idx_source_project (source_project_id),
                INDEX idx_target_project (target_project_id),
                INDEX idx_assigner (assigner_id),
                INDEX idx_assignee (assignee_id),
                INDEX idx_status (status)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='协同任务分配表'
        """)
        
        # 创建项目同步日志表
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `project_sync_logs` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `collaboration_task_id` INT NOT NULL COMMENT '协同任务ID',
                `sync_type` ENUM('asset_add', 'asset_update', 'asset_delete') NOT NULL COMMENT '同步类型',
                `asset_id` INT DEFAULT NULL COMMENT '资产ID',
                `sync_data` TEXT DEFAULT NULL COMMENT '同步数据（JSON）',
                `sync_direction` ENUM('to_source', 'to_target') NOT NULL COMMENT '同步方向',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (collaboration_task_id) REFERENCES collaboration_tasks(id) ON DELETE CASCADE,
                INDEX idx_task_id (collaboration_task_id),
                INDEX idx_sync_type (sync_type),
                INDEX idx_created_at (created_at)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='项目同步日志表'
        """)
        
        # 创建项目访问权限表（细粒度权限控制）
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS `project_access_permissions` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `project_id` INT NOT NULL COMMENT '项目ID',
                `user_id` INT NOT NULL COMMENT '用户ID',
                `can_view` BOOLEAN DEFAULT TRUE COMMENT '是否可查看',
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP,
                `updated_at` DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                FOREIGN KEY (project_id) REFERENCES projects(id) ON DELETE CASCADE,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
                UNIQUE KEY unique_permission (project_id, user_id),
                INDEX idx_project_id (project_id),
                INDEX idx_user_id (user_id)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
            COMMENT='项目访问权限表'
        """)
        
        # 创建默认超级管理员账户
        default_username = 'admin'
        default_password = 'admin123'
        password_hash = hashlib.sha256(default_password.encode()).hexdigest()
        
        cursor.execute("""
            INSERT INTO users (username, password, is_super_admin) 
            VALUES (%s, %s, TRUE)
        """, (default_username, password_hash))
        
        connection.commit()
        cursor.close()
        connection.close()
        
        # 保存配置
        config = {
            'host': host,
            'port': int(port),
            'user': user,
            'password': password,
            'database': database
        }
        save_config(config)
        
        return jsonify({'success': True, 'message': '安装成功！默认管理员账户: admin/admin123'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'安装失败: {str(e)}'}), 500

@app.route('/api/login', methods=['POST'])
def login():
    """用户登录"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装，请先安装'}), 400
    
    data = request.json
    username = data.get('username')
    password = data.get('password')
    
    if not username or not password:
        return jsonify({'success': False, 'message': '用户名和密码不能为空'}), 400
    
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, username, is_admin, is_super_admin FROM users WHERE username = %s AND password = %s",
            (username, password_hash)
        )
        user = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if user:
            # 生成JWT token
            token = generate_token(
                user['id'],
                user['username'],
                user.get('is_admin', False),
                user.get('is_super_admin', False)
            )
            
            return jsonify({
                'success': True,
                'message': '登录成功',
                'token': token,  # 新增：返回JWT token
                'user': {
                    'id': user['id'],
                    'username': user['username'],
                    'is_admin': user.get('is_admin', False),
                    'is_super_admin': user.get('is_super_admin', False)
                }
            })
        else:
            return jsonify({'success': False, 'message': '用户名或密码错误'}), 401
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'登录失败: {str(e)}'}), 500

@app.route('/api/register', methods=['POST'])
def register():
    """用户注册（支持邀请码）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装，请先安装'}), 400
    
    data = request.json
    username = data.get('username')
    password = data.get('password')
    invite_code = data.get('invite_code')  # 可选的邀请码
    
    if not username or not password:
        return jsonify({'success': False, 'message': '用户名和密码不能为空'}), 400
    
    if len(username) < 3 or len(username) > 50:
        return jsonify({'success': False, 'message': '用户名长度必须在3-50个字符之间'}), 400
    
    if len(password) < 6:
        return jsonify({'success': False, 'message': '密码长度至少为6个字符'}), 400
    
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查是否允许注册
        cursor.execute("""
            SELECT setting_value 
            FROM system_settings 
            WHERE setting_key = 'allow_registration'
        """)
        result = cursor.fetchone()
        
        # 如果设置为不允许注册，则拒绝（除非有邀请码）
        if result and result['setting_value'].lower() == 'false' and not invite_code:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '系统当前不允许注册，请使用邀请码注册'}), 403
        
        inviter_id = None
        
        # 如果提供了邀请码，验证并获取邀请人ID
        if invite_code:
            cursor.execute(
                "SELECT user_id, is_used FROM invite_codes WHERE code = %s",
                (invite_code,)
            )
            invite_info = cursor.fetchone()
            
            if not invite_info:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '邀请码无效'}), 400
            
            if invite_info['is_used']:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '邀请码已被使用'}), 400
            
            inviter_id = invite_info['user_id']
            
            # 标记邀请码为已使用
            cursor.execute(
                "UPDATE invite_codes SET is_used = TRUE WHERE code = %s",
                (invite_code,)
            )
        
        # 创建用户
        cursor.execute(
            "INSERT INTO users (username, password) VALUES (%s, %s)",
            (username, password_hash)
        )
        user_id = cursor.lastrowid
        
        # 如果是通过邀请码注册，建立协作关系
        if inviter_id:
            # 双向协作关系
            cursor.execute(
                "INSERT INTO user_collaborations (user_id, collaborator_id) VALUES (%s, %s)",
                (inviter_id, user_id)
            )
            cursor.execute(
                "INSERT INTO user_collaborations (user_id, collaborator_id) VALUES (%s, %s)",
                (user_id, inviter_id)
            )
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '注册成功' + ('，已加入协作团队' if inviter_id else ''),
            'user': {
                'id': user_id,
                'username': username
            }
        })
    
    except pymysql.IntegrityError:
        return jsonify({'success': False, 'message': '用户名已存在'}), 400
    except Exception as e:
        return jsonify({'success': False, 'message': f'注册失败: {str(e)}'}), 500

# ==================== FOFA凭证管理 ====================

@app.route('/api/user/fofa-credentials', methods=['GET'])
def get_fofa_credentials():
    """获取当前用户的FOFA凭证"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT fofa_email, fofa_api_key FROM users WHERE id = %s",
            (user_id,)
        )
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if result:
            return jsonify({
                'success': True,
                'fofa_email': result['fofa_email'] or '',
                'fofa_api_key': result['fofa_api_key'] or ''
            })
        else:
            return jsonify({'success': False, 'message': '用户不存在'}), 404
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取FOFA凭证失败: {str(e)}'}), 500

@app.route('/api/user/fofa-credentials', methods=['PUT'])
def update_fofa_credentials():
    """更新当前用户的FOFA凭证"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    fofa_email = data.get('fofa_email', '').strip()
    fofa_api_key = data.get('fofa_api_key', '').strip()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE users SET fofa_email = %s, fofa_api_key = %s WHERE id = %s",
            (fofa_email if fofa_email else None, fofa_api_key if fofa_api_key else None, user_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': 'FOFA凭证保存成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'保存FOFA凭证失败: {str(e)}'}), 500

# ==================== 报告模板管理 ====================

@app.route('/api/report-templates', methods=['GET'])
def get_report_templates():
    """获取报告模板列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT t.*, u.username as creator_name
            FROM report_templates t
            LEFT JOIN users u ON t.created_by = u.id
            WHERE t.created_by = %s
            ORDER BY t.created_at DESC
        """, (user_id,))
        templates = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'templates': templates
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取模板列表失败: {str(e)}'}), 500

@app.route('/api/report-templates', methods=['POST'])
def upload_report_template():
    """上传报告模板"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    # 检查是否有文件
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    # 检查文件类型
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'message': '只支持 .docx 格式的文件'}), 400
    
    # 获取表单数据
    name = request.form.get('name', '')
    description = request.form.get('description', '')
    
    if not name:
        return jsonify({'success': False, 'message': '模板名称不能为空'}), 400
    
    try:
        # 创建上传目录
        upload_dir = 'uploads/report_templates'
        os.makedirs(upload_dir, exist_ok=True)
        
        # 生成唯一文件名
        import time
        timestamp = int(time.time() * 1000)
        file_ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{user_id}_{timestamp}{file_ext}"
        file_path = os.path.join(upload_dir, unique_filename)
        
        # 保存文件
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # 解析 docx 文件，提取占位符
        placeholders = []
        try:
            from docx import Document
            doc = Document(file_path)
            
            # 提取文档中的所有文本
            full_text = []
            
            # 提取正文段落
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            # 提取页眉和页脚
            for section in doc.sections:
                # 提取页眉
                header = section.header
                for para in header.paragraphs:
                    full_text.append(para.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
                
                # 提取页脚
                footer = section.footer
                for para in footer.paragraphs:
                    full_text.append(para.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)
            
            # 使用正则表达式提取 {{占位符}}
            import re
            text_content = '\n'.join(full_text)
            placeholder_pattern = r'\{\{([^}]+)\}\}'
            found_placeholders = re.findall(placeholder_pattern, text_content)
            
            # 去重并保存
            placeholders = list(set(found_placeholders))
            print(f"从模板中提取到的占位符: {placeholders}")
        except Exception as e:
            print(f"解析占位符失败: {str(e)}")
            import traceback
            traceback.print_exc()
            # 即使解析失败也继续保存模板
        
        placeholders_json = json.dumps(placeholders, ensure_ascii=False) if placeholders else None
        
        # 保存到数据库
        conn = get_db_connection()
        if not conn:
            os.remove(file_path)  # 删除已上传的文件
            return jsonify({'success': False, 'message': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO report_templates (name, description, file_name, file_path, file_size, placeholders, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (name, description, file.filename, file_path, file_size, placeholders_json, user_id))
        
        template_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '模板上传成功',
            'template_id': template_id
        })
    
    except Exception as e:
        # 如果出错，删除已上传的文件
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({'success': False, 'message': f'上传失败: {str(e)}'}), 500

@app.route('/api/report-templates/<int:template_id>', methods=['DELETE'])
def delete_report_template(template_id):
    """删除报告模板"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查模板是否存在且属于当前用户
        cursor.execute("""
            SELECT file_path FROM report_templates 
            WHERE id = %s AND created_by = %s
        """, (template_id, user_id))
        template = cursor.fetchone()
        
        if not template:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '模板不存在或无权删除'}), 404
        
        # 删除文件
        file_path = template['file_path']
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # 删除数据库记录
        cursor.execute("DELETE FROM report_templates WHERE id = %s", (template_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '模板删除成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除失败: {str(e)}'}), 500

@app.route('/api/report-templates/<int:template_id>/download', methods=['GET'])
def download_report_template(template_id):
    """下载报告模板"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT file_path, file_name FROM report_templates 
            WHERE id = %s AND created_by = %s
        """, (template_id, user_id))
        template = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not template:
            return jsonify({'success': False, 'message': '模板不存在或无权访问'}), 404
        
        file_path = template['file_path']
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'message': '文件不存在'}), 404
        
        return send_file(file_path, as_attachment=True, download_name=template['file_name'])
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'下载失败: {str(e)}'}), 500

# ==================== 报告管理 ====================

@app.route('/api/reports', methods=['GET'])
def get_reports():
    """获取报告列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    project_id = request.args.get('project_id')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        if project_id:
            cursor.execute("""
                SELECT r.*, t.name as template_name, u.username as creator_name
                FROM reports r
                LEFT JOIN report_templates t ON r.template_id = t.id
                LEFT JOIN users u ON r.created_by = u.id
                WHERE r.created_by = %s AND r.project_id = %s
                ORDER BY r.updated_at DESC
            """, (user_id, project_id))
        else:
            cursor.execute("""
                SELECT r.*, t.name as template_name, u.username as creator_name
                FROM reports r
                LEFT JOIN report_templates t ON r.template_id = t.id
                LEFT JOIN users u ON r.created_by = u.id
                WHERE r.created_by = %s
                ORDER BY r.updated_at DESC
            """, (user_id,))
        
        reports = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'reports': reports
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取报告列表失败: {str(e)}'}), 500

@app.route('/api/reports', methods=['POST'])
def create_report():
    """创建报告"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    template_id = data.get('template_id')
    name = data.get('name')
    project_id = data.get('project_id')
    
    if not template_id or not name:
        return jsonify({'success': False, 'message': '模板ID和报告名称不能为空'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取模板信息
        cursor.execute("""
            SELECT placeholders FROM report_templates 
            WHERE id = %s AND created_by = %s
        """, (template_id, user_id))
        template = cursor.fetchone()
        
        if not template:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '模板不存在或无权访问'}), 404
        
        # 初始化报告内容
        placeholders = json.loads(template['placeholders']) if template['placeholders'] else []
        initial_content = {
            'placeholders': {p: '' for p in placeholders},
            'vulnerabilities': []
        }
        
        cursor.execute("""
            INSERT INTO reports (template_id, project_id, name, content, created_by)
            VALUES (%s, %s, %s, %s, %s)
        """, (template_id, project_id, name, json.dumps(initial_content, ensure_ascii=False), user_id))
        
        report_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '报告创建成功',
            'report_id': report_id
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建报告失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>', methods=['GET'])
def get_report(report_id):
    """获取报告详情"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT r.*, t.name as template_name, t.placeholders
            FROM reports r
            LEFT JOIN report_templates t ON r.template_id = t.id
            WHERE r.id = %s AND r.created_by = %s
        """, (report_id, user_id))
        report = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not report:
            return jsonify({'success': False, 'message': '报告不存在或无权访问'}), 404
        
        return jsonify({
            'success': True,
            'report': report
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取报告失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>', methods=['PUT'])
def update_report(report_id):
    """更新报告内容"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    content = data.get('content')
    status = data.get('status')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查报告是否存在
        cursor.execute("""
            SELECT id FROM reports 
            WHERE id = %s AND created_by = %s
        """, (report_id, user_id))
        
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '报告不存在或无权访问'}), 404
        
        # 更新报告
        update_fields = []
        update_values = []
        
        if content is not None:
            update_fields.append('content = %s')
            update_values.append(json.dumps(content, ensure_ascii=False))
        
        if status:
            update_fields.append('status = %s')
            update_values.append(status)
        
        if not update_fields:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '没有要更新的内容'}), 400
        
        update_values.append(report_id)
        
        cursor.execute(f"""
            UPDATE reports 
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, update_values)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '报告更新成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新报告失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>', methods=['DELETE'])
def delete_report(report_id):
    """删除报告"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查报告是否存在
        cursor.execute("""
            SELECT id FROM reports 
            WHERE id = %s AND created_by = %s
        """, (report_id, user_id))
        
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '报告不存在或无权访问'}), 404
        
        # 删除报告（附件会通过外键级联删除）
        cursor.execute("DELETE FROM reports WHERE id = %s", (report_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '报告删除成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除报告失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>/upload-image', methods=['POST'])
def upload_report_image(report_id):
    """上传报告图片"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400
    
    # 检查文件类型
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}
    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    
    if file_ext not in allowed_extensions:
        return jsonify({'success': False, 'message': '只支持图片格式'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查报告是否存在
        cursor.execute("""
            SELECT id FROM reports 
            WHERE id = %s AND created_by = %s
        """, (report_id, user_id))
        
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '报告不存在或无权访问'}), 404
        
        # 创建上传目录
        upload_dir = f'uploads/report_images/{report_id}'
        os.makedirs(upload_dir, exist_ok=True)
        
        # 生成唯一文件名
        import time
        timestamp = int(time.time() * 1000)
        unique_filename = f"{timestamp}.{file_ext}"
        file_path = os.path.join(upload_dir, unique_filename)
        
        # 保存文件
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # 保存到数据库
        cursor.execute("""
            INSERT INTO report_attachments (report_id, file_name, file_path, file_type, file_size)
            VALUES (%s, %s, %s, %s, %s)
        """, (report_id, file.filename, file_path, file_ext, file_size))
        
        attachment_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '图片上传成功',
            'attachment_id': attachment_id,
            'url': f'/api/reports/{report_id}/images/{attachment_id}'
        })
    
    except Exception as e:
        if 'file_path' in locals() and os.path.exists(file_path):
            os.remove(file_path)
        return jsonify({'success': False, 'message': f'上传失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>/images/<int:attachment_id>', methods=['GET'])
def get_report_image(report_id, attachment_id):
    """获取报告图片"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT a.file_path, a.file_name
            FROM report_attachments a
            JOIN reports r ON a.report_id = r.id
            WHERE a.id = %s AND a.report_id = %s AND r.created_by = %s
        """, (attachment_id, report_id, user_id))
        
        attachment = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not attachment:
            return jsonify({'success': False, 'message': '图片不存在或无权访问'}), 404
        
        file_path = attachment['file_path']
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'message': '文件不存在'}), 404
        
        return send_file(file_path, mimetype=f'image/{os.path.splitext(file_path)[1][1:]}')
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取图片失败: {str(e)}'}), 500

@app.route('/api/reports/<int:report_id>/export', methods=['GET'])
def export_report(report_id):
    """导出报告为 Word 文档"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        cursor = conn.cursor()
        
        # 获取报告信息
        cursor.execute("""
            SELECT r.*, t.file_path as template_path, t.name as template_name
            FROM reports r
            JOIN report_templates t ON r.template_id = t.id
            WHERE r.id = %s AND r.created_by = %s
        """, (report_id, user_id))
        
        report = cursor.fetchone()
        
        if not report:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '报告不存在或无权访问'}), 404
        
        # 解析报告内容
        content = json.loads(report['content']) if report['content'] else {'placeholders': {}, 'vulnerabilities': []}
        
        # 打印调试信息
        print(f"报告内容: {content}")
        print(f"占位符: {content.get('placeholders', {})}")
        
        # 加载模板文档
        template_path = report['template_path']
        if not os.path.exists(template_path):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '模板文件不存在'}), 404
        
        doc = Document(template_path)
        
        # 1. 替换占位符
        placeholders = content.get('placeholders', {})
        print(f"准备替换的占位符: {placeholders}")
        
        def replace_placeholder_in_paragraph(paragraph, placeholders):
            """在段落中替换占位符，处理跨 run 的情况"""
            full_text = paragraph.text
            
            # 检查是否包含占位符
            has_placeholder = False
            for key in placeholders.keys():
                if f'{{{{{key}}}}}' in full_text:
                    has_placeholder = True
                    print(f"找到占位符 {{{{{key}}}}} 在段落: {full_text[:50]}...")
                    break
            
            if not has_placeholder:
                return
            
            # 替换所有占位符
            new_text = full_text
            for key, value in placeholders.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in new_text:
                    print(f"替换 {placeholder} -> {value}")
                    new_text = new_text.replace(placeholder, str(value))
            
            # 如果文本有变化，重新设置段落内容
            if new_text != full_text:
                print(f"段落文本已更新: {full_text[:50]}... -> {new_text[:50]}...")
                # 保存第一个 run 的格式
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    font_bold = first_run.font.bold
                    font_italic = first_run.font.italic
                    
                    # 清空所有 runs
                    for run in paragraph.runs:
                        run.text = ''
                    
                    # 设置新文本到第一个 run
                    paragraph.runs[0].text = new_text
                    paragraph.runs[0].font.name = font_name
                    paragraph.runs[0].font.size = font_size
                    paragraph.runs[0].font.bold = font_bold
                    paragraph.runs[0].font.italic = font_italic
                else:
                    # 如果没有 run，添加一个
                    paragraph.add_run(new_text)
        
        # 2. 准备漏洞列表内容
        vulnerabilities = content.get('vulnerabilities', [])
        
        def generate_vulnerability_content():
            """生成漏洞列表的文本内容"""
            if not vulnerabilities:
                return "暂无漏洞信息"
            
            vuln_text = []
            for idx, vuln in enumerate(vulnerabilities, 1):
                vuln_text.append(f"\n{idx}. {vuln.get('name', '未命名漏洞')}\n")
                
                # 危害等级
                severity_map = {
                    'critical': '严重',
                    'high': '高危',
                    'medium': '中危',
                    'low': '低危',
                    'info': '信息'
                }
                severity_text = severity_map.get(vuln.get('severity', 'medium'), '中危')
                vuln_text.append(f"危害等级：{severity_text}\n")
                
                # 漏洞描述
                if vuln.get('description'):
                    vuln_text.append(f"\n漏洞描述：\n{vuln['description']}\n")
                
                # 复现步骤
                if vuln.get('reproduction'):
                    vuln_text.append(f"\n复现步骤：\n{vuln['reproduction']}\n")
                
                # 修复建议
                if vuln.get('solution'):
                    vuln_text.append(f"\n修复建议：\n{vuln['solution']}\n")
                
                if idx < len(vulnerabilities):
                    vuln_text.append("\n" + "─" * 50 + "\n")
            
            return ''.join(vuln_text)
        
        # 将漏洞信息作为特殊占位符
        vuln_content = generate_vulnerability_content()
        placeholders['漏洞信息'] = vuln_content
        
        # 检查是否有 {{漏洞信息}} 占位符
        has_vuln_placeholder = False
        
        # 替换段落中的占位符（包括漏洞信息）
        for paragraph in doc.paragraphs:
            if '{{漏洞信息}}' in paragraph.text:
                has_vuln_placeholder = True
            replace_placeholder_in_paragraph(paragraph, placeholders)
        
        # 替换表格中的占位符（包括漏洞信息）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{漏洞信息}}' in paragraph.text:
                            has_vuln_placeholder = True
                        replace_placeholder_in_paragraph(paragraph, placeholders)
        
        # 替换页眉中的占位符
        for section in doc.sections:
            # 替换页眉
            header = section.header
            for paragraph in header.paragraphs:
                if '{{漏洞信息}}' in paragraph.text:
                    has_vuln_placeholder = True
                replace_placeholder_in_paragraph(paragraph, placeholders)
            
            # 替换页眉中的表格
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{漏洞信息}}' in paragraph.text:
                                has_vuln_placeholder = True
                            replace_placeholder_in_paragraph(paragraph, placeholders)
            
            # 替换页脚
            footer = section.footer
            for paragraph in footer.paragraphs:
                if '{{漏洞信息}}' in paragraph.text:
                    has_vuln_placeholder = True
                replace_placeholder_in_paragraph(paragraph, placeholders)
            
            # 替换页脚中的表格
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{漏洞信息}}' in paragraph.text:
                                has_vuln_placeholder = True
                            replace_placeholder_in_paragraph(paragraph, placeholders)
        
        # 如果没有使用 {{漏洞信息}} 占位符，则在文档末尾添加详细的漏洞列表
        if vulnerabilities and not has_vuln_placeholder:
            # 在文档末尾添加漏洞详情
            doc.add_page_break()
            
            # 添加标题
            heading = doc.add_heading('漏洞详情', level=1)
            
            for idx, vuln in enumerate(vulnerabilities, 1):
                # 漏洞标题
                vuln_heading = doc.add_heading(f'{idx}. {vuln.get("name", "未命名漏洞")}', level=2)
                
                # 危害等级
                severity_map = {
                    'critical': '严重',
                    'high': '高危',
                    'medium': '中危',
                    'low': '低危',
                    'info': '信息'
                }
                severity_text = severity_map.get(vuln.get('severity', 'medium'), '中危')
                severity_para = doc.add_paragraph()
                severity_para.add_run('危害等级：').bold = True
                severity_run = severity_para.add_run(severity_text)
                
                # 根据等级设置颜色
                if vuln.get('severity') == 'critical':
                    severity_run.font.color.rgb = RGBColor(139, 0, 0)  # 深红色
                elif vuln.get('severity') == 'high':
                    severity_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
                elif vuln.get('severity') == 'medium':
                    severity_run.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
                elif vuln.get('severity') == 'low':
                    severity_run.font.color.rgb = RGBColor(255, 255, 0)  # 黄色
                
                # 漏洞描述
                if vuln.get('description'):
                    desc_para = doc.add_paragraph()
                    desc_para.add_run('漏洞描述：').bold = True
                    doc.add_paragraph(vuln['description'])
                
                # 复现步骤
                if vuln.get('reproduction'):
                    repro_para = doc.add_paragraph()
                    repro_para.add_run('复现步骤：').bold = True
                    doc.add_paragraph(vuln['reproduction'])
                
                # 修复建议
                if vuln.get('solution'):
                    solution_para = doc.add_paragraph()
                    solution_para.add_run('修复建议：').bold = True
                    doc.add_paragraph(vuln['solution'])
                
                # 插入截图
                images = vuln.get('images', [])
                if images:
                    img_para = doc.add_paragraph()
                    img_para.add_run('相关截图：').bold = True
                    
                    for img in images:
                        # 获取图片路径
                        cursor.execute("""
                            SELECT file_path FROM report_attachments
                            WHERE id = %s AND report_id = %s
                        """, (img['name'], report_id))
                        
                        img_record = cursor.fetchone()
                        if img_record and os.path.exists(img_record['file_path']):
                            try:
                                doc.add_picture(img_record['file_path'], width=Inches(5))
                                doc.add_paragraph()  # 添加空行
                            except Exception as e:
                                print(f"插入图片失败: {str(e)}")
                
                # 添加分隔线
                if idx < len(vulnerabilities):
                    doc.add_paragraph('─' * 50)
        
        # 3. 保存导出的文档
        export_dir = 'uploads/exports'
        os.makedirs(export_dir, exist_ok=True)
        
        import time
        timestamp = int(time.time() * 1000)
        export_filename = f"report_{report_id}_{timestamp}.docx"
        export_path = os.path.join(export_dir, export_filename)
        
        doc.save(export_path)
        
        cursor.close()
        conn.close()
        
        # 返回文件
        return send_file(
            export_path,
            as_attachment=True,
            download_name=f"{report['name']}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        print(f"导出报告失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'导出失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>', methods=['GET'])
def get_user(user_id):
    """获取用户信息"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能查询自己的信息，除非是管理员
    if current_user['id'] != user_id and not current_user.get('is_admin') and not current_user.get('is_super_admin'):
        return jsonify({'success': False, 'message': '无权访问该用户信息'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, username, created_at FROM users WHERE id = %s",
            (user_id,)
        )
        user = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if user:
            return jsonify({
                'success': True,
                'user': user
            })
        else:
            return jsonify({'success': False, 'message': '用户不存在'}), 404
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取用户信息失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/username', methods=['PUT'])
def update_username(user_id):
    """修改用户昵称"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能修改自己的用户名
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权修改该用户的用户名'}), 403
    
    data = request.json
    new_username = data.get('username')
    
    if not new_username:
        return jsonify({'success': False, 'message': '用户名不能为空'}), 400
    
    if len(new_username) < 3 or len(new_username) > 50:
        return jsonify({'success': False, 'message': '用户名长度必须在3-50个字符之间'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查用户名是否已存在
        cursor.execute(
            "SELECT id FROM users WHERE username = %s AND id != %s",
            (new_username, user_id)
        )
        if cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '用户名已被使用'}), 400
        
        # 更新用户名
        cursor.execute(
            "UPDATE users SET username = %s WHERE id = %s",
            (new_username, user_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '用户名修改成功',
            'username': new_username
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'修改用户名失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/password', methods=['PUT'])
def update_password(user_id):
    """修改用户密码"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能修改自己的密码
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权修改该用户的密码'}), 403
    
    data = request.json
    old_password = data.get('old_password')
    new_password = data.get('new_password')
    
    if not old_password or not new_password:
        return jsonify({'success': False, 'message': '旧密码和新密码不能为空'}), 400
    
    if len(new_password) < 6:
        return jsonify({'success': False, 'message': '新密码长度至少为6个字符'}), 400
    
    old_password_hash = hashlib.sha256(old_password.encode()).hexdigest()
    new_password_hash = hashlib.sha256(new_password.encode()).hexdigest()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 验证旧密码
        cursor.execute(
            "SELECT id FROM users WHERE id = %s AND password = %s",
            (user_id, old_password_hash)
        )
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '旧密码错误'}), 401
        
        # 更新密码
        cursor.execute(
            "UPDATE users SET password = %s WHERE id = %s",
            (new_password_hash, user_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '密码修改成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'修改密码失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/invite-code', methods=['POST'])
def generate_invite_code(user_id):
    """生成邀请码"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能为自己生成邀请码
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权为该用户生成邀请码'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 生成唯一的邀请码
        code = str(uuid.uuid4()).replace('-', '')[:16]
        
        # 设置过期时间（7天后）
        expires_at = datetime.now() + timedelta(days=7)
        
        cursor.execute(
            "INSERT INTO invite_codes (code, user_id, expires_at) VALUES (%s, %s, %s)",
            (code, user_id, expires_at)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '邀请码生成成功',
            'invite_code': code,
            'expires_at': expires_at.strftime('%Y-%m-%d %H:%M:%S')
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'生成邀请码失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/invite-codes', methods=['GET'])
def get_invite_codes(user_id):
    """获取用户的邀请码列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能查看自己的邀请码
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权查看该用户的邀请码'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT code, created_at, expires_at, is_used
            FROM invite_codes
            WHERE user_id = %s
            ORDER BY created_at DESC
        """, (user_id,))
        codes = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'invite_codes': codes
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取邀请码失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/collaborators', methods=['GET'])
def get_collaborators(user_id):
    """获取协作者列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能查看自己的协作者列表
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权查看该用户的协作者列表'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT u.id, u.username, uc.created_at
            FROM user_collaborations uc
            JOIN users u ON uc.collaborator_id = u.id
            WHERE uc.user_id = %s
            ORDER BY uc.created_at DESC
        """, (user_id,))
        collaborators = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'collaborators': collaborators
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取协作者失败: {str(e)}'}), 500

@app.route('/api/users/<int:user_id>/collaborators/<int:collaborator_id>', methods=['DELETE'])
def delete_collaborator(user_id, collaborator_id):
    """删除协作者（双向删除协作关系）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    current_user = get_current_user()
    if not current_user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 权限检查：用户只能删除自己的协作者
    if current_user['id'] != user_id:
        return jsonify({'success': False, 'message': '无权删除该用户的协作者'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 双向删除协作关系
        # 删除 user_id -> collaborator_id 的关系
        cursor.execute("""
            DELETE FROM user_collaborations 
            WHERE user_id = %s AND collaborator_id = %s
        """, (user_id, collaborator_id))
        
        # 删除 collaborator_id -> user_id 的关系
        cursor.execute("""
            DELETE FROM user_collaborations 
            WHERE user_id = %s AND collaborator_id = %s
        """, (collaborator_id, user_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '协作关系已删除'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除协作者失败: {str(e)}'}), 500

@app.route('/api/projects', methods=['GET'])
def get_projects():
    """获取项目列表（支持协作权限）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查 is_visible_to_inviter 字段是否存在
        cursor.execute("SHOW COLUMNS FROM projects LIKE 'is_visible_to_inviter'")
        has_visibility_field = cursor.fetchone() is not None
        
        if has_visibility_field:
            # 检查权限表是否存在
            cursor.execute("SHOW TABLES LIKE 'project_access_permissions'")
            has_permission_table = cursor.fetchone() is not None
            
            if has_permission_table:
                # 使用新的查询（支持细粒度权限控制）
                cursor.execute("""
                    SELECT DISTINCT p.*, u.username as creator_name,
                           (SELECT COUNT(*) FROM assets WHERE project_id = p.id) as asset_count,
                           assigner.username as assigner_name
                    FROM projects p
                    LEFT JOIN users u ON p.created_by = u.id
                    LEFT JOIN users assigner ON p.task_assigner_id = assigner.id
                    LEFT JOIN project_access_permissions pap ON p.id = pap.project_id AND pap.user_id = %s
                    WHERE p.created_by = %s
                       OR (p.created_by IN (
                               SELECT collaborator_id FROM user_collaborations WHERE user_id = %s
                           ) 
                           AND (p.is_visible_to_inviter = TRUE OR p.is_visible_to_inviter IS NULL)
                           AND (pap.can_view IS NULL OR pap.can_view = TRUE))
                    ORDER BY p.updated_at DESC
                """, (user_id, user_id, user_id))
            else:
                # 使用旧的查询（支持可见性控制）
                cursor.execute("""
                    SELECT DISTINCT p.*, u.username as creator_name,
                           (SELECT COUNT(*) FROM assets WHERE project_id = p.id) as asset_count,
                           assigner.username as assigner_name
                    FROM projects p
                    LEFT JOIN users u ON p.created_by = u.id
                    LEFT JOIN users assigner ON p.task_assigner_id = assigner.id
                    WHERE p.created_by = %s
                       OR (p.created_by IN (
                               SELECT collaborator_id FROM user_collaborations WHERE user_id = %s
                           ) AND (p.is_visible_to_inviter = TRUE OR p.is_visible_to_inviter IS NULL))
                    ORDER BY p.updated_at DESC
                """, (user_id, user_id))
        else:
            # 使用旧的查询（兼容模式）
            cursor.execute("""
                SELECT DISTINCT p.*, u.username as creator_name,
                       (SELECT COUNT(*) FROM assets WHERE project_id = p.id) as asset_count
                FROM projects p
                LEFT JOIN users u ON p.created_by = u.id
                WHERE p.created_by = %s
                   OR p.created_by IN (
                       SELECT collaborator_id FROM user_collaborations WHERE user_id = %s
                   )
                ORDER BY p.updated_at DESC
            """, (user_id, user_id))
        
        projects = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'projects': projects})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取项目列表失败: {str(e)}'}), 500

@app.route('/api/projects', methods=['POST'])
def create_project():
    """创建新项目"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    name = data.get('name')
    description = data.get('description', '')
    
    if not name:
        return jsonify({'success': False, 'message': '项目名称不能为空'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO projects (name, description, created_by) VALUES (%s, %s, %s)",
            (name, description, user_id)
        )
        conn.commit()
        project_id = cursor.lastrowid
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '项目创建成功',
            'project_id': project_id
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建项目失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>', methods=['PUT'])
def update_project(project_id):
    """更新项目信息"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    name = data.get('name')
    description = data.get('description')
    
    if not name:
        return jsonify({'success': False, 'message': '项目名称不能为空'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：只有项目所有者才能更新
        if not check_project_permission(cursor, project_id, user_id, require_owner=True):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此项目'}), 403
        
        cursor.execute(
            "UPDATE projects SET name = %s, description = %s WHERE id = %s",
            (name, description, project_id)
        )
        conn.commit()
        
        if cursor.rowcount == 0:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '项目不存在'}), 404
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '项目更新成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新项目失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    """删除项目"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：只有项目所有者才能删除
        if not check_project_permission(cursor, project_id, user_id, require_owner=True):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此项目'}), 403
        
        cursor.execute("DELETE FROM projects WHERE id = %s", (project_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '项目删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除项目失败: {str(e)}'}), 500

def parse_asset(asset_str):
    """解析资产字符串，识别类型 - 增强版支持多种格式"""
    asset_str = asset_str.strip()
    
    if not asset_str:
        return None
    
    # 尝试从JSON或类JSON结构中提取信息
    try:
        # 尝试修复常见的JSON格式问题
        fixed_str = asset_str
        
        # 移除开头的冒号和逗号
        fixed_str = re.sub(r'^[,:]+', '', fixed_str)
        
        # 如果不是以{开头，尝试添加
        if not fixed_str.startswith('{'):
            fixed_str = '{' + fixed_str
        
        # 如果不是以}结尾，尝试添加
        if not fixed_str.endswith('}'):
            fixed_str = fixed_str + '}'
        
        # 尝试解析JSON
        try:
            data = json.loads(fixed_str)
            
            # 从JSON中提取URL
            url = data.get('url')
            ip = data.get('ip')
            
            if url:
                # 解析URL
                protocol = 'https' if url.startswith('https://') else 'http'
                match = re.match(r'(https?://)?([^:/]+)(?::(\d+))?', url)
                if match:
                    host = match.group(2)
                    port = int(match.group(3)) if match.group(3) else (443 if protocol == 'https' else 80)
                    
                    # 如果URL中没有协议，添加它
                    if not url.startswith('http'):
                        url = f"{protocol}://{url}"
                    
                    return {
                        'asset_type': protocol,
                        'asset_value': url,
                        'ip': ip if ip else host,
                        'port': port,
                        'protocol': protocol,
                        'url': url
                    }
        except (json.JSONDecodeError, ValueError):
            pass
    except Exception:
        pass
    
    # 使用正则表达式提取URL（从任意文本中）
    url_pattern = r'(https?://[^\s\'"<>{}]+(?::\d+)?(?:/[^\s\'"<>{}]*)?)'
    url_matches = re.findall(url_pattern, asset_str)
    if url_matches:
        url = url_matches[0]
        protocol = 'https' if url.startswith('https://') else 'http'
        match = re.match(r'https?://([^:/]+)(?::(\d+))?', url)
        if match:
            host = match.group(1)
            port = int(match.group(2)) if match.group(2) else (443 if protocol == 'https' else 80)
            
            # 尝试从原始字符串中提取IP
            ip_pattern = r'"ip"\s*:\s*"([^"]+)"'
            ip_match = re.search(ip_pattern, asset_str)
            ip = ip_match.group(1) if ip_match else host
            
            return {
                'asset_type': protocol,
                'asset_value': url,
                'ip': ip,
                'port': port,
                'protocol': protocol,
                'url': url
            }
    
    # 提取域名并构造URL（从host字段）
    host_pattern = r'"host"\s*:\s*"([^"]+)"'
    host_match = re.search(host_pattern, asset_str)
    if host_match:
        host = host_match.group(1)
        
        # 提取IP
        ip_pattern = r'"ip"\s*:\s*"([^"]+)"'
        ip_match = re.search(ip_pattern, asset_str)
        ip = ip_match.group(1) if ip_match else host
        
        # 判断协议
        protocol = 'https' if 'ssl' in asset_str.lower() or 'https' in asset_str.lower() else 'http'
        
        # 提取端口（从URL或其他地方）
        port_pattern = r':(\d+)'
        port_matches = re.findall(port_pattern, asset_str)
        port = int(port_matches[-1]) if port_matches else (443 if protocol == 'https' else 80)
        
        # 提取路径
        path_pattern = r'"path"\s*:\s*"([^"]*)"'
        path_match = re.search(path_pattern, asset_str)
        path = path_match.group(1) if path_match else '/'
        
        # 构造URL
        url = f"{protocol}://{host}:{port}{path}" if path else f"{protocol}://{host}:{port}"
        
        return {
            'asset_type': protocol,
            'asset_value': url,
            'ip': ip,
            'port': port,
            'protocol': protocol,
            'url': url
        }
    
    # 标准格式：完整URL
    if asset_str.startswith('http://') or asset_str.startswith('https://'):
        protocol = 'https' if asset_str.startswith('https://') else 'http'
        match = re.match(r'https?://([^:/]+)(?::(\d+))?', asset_str)
        if match:
            ip = match.group(1)
            port = int(match.group(2)) if match.group(2) else (443 if protocol == 'https' else 80)
            return {
                'asset_type': protocol,
                'asset_value': asset_str,
                'ip': ip,
                'port': port,
                'protocol': protocol,
                'url': asset_str
            }
    
    # 匹配 IP:Port
    match = re.match(r'^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}):(\d+)$', asset_str)
    if match:
        ip = match.group(1)
        port = int(match.group(2))
        return {
            'asset_type': 'ip_port',
            'asset_value': asset_str,
            'ip': ip,
            'port': port,
            'protocol': None,
            'url': None
        }
    
    # 匹配纯IP
    match = re.match(r'^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})$', asset_str)
    if match:
        ip = match.group(1)
        return {
            'asset_type': 'ip',
            'asset_value': asset_str,
            'ip': ip,
            'port': None,
            'protocol': None,
            'url': None
        }
    
    # 匹配域名（简单的域名格式）- 保持原样，不自动转换为http
    # 排除文件名（带有常见文件扩展名的）
    common_file_extensions = [
        'js', 'css', 'html', 'htm', 'xml', 'json', 'txt', 'pdf', 'doc', 'docx',
        'xls', 'xlsx', 'ppt', 'pptx', 'zip', 'rar', 'tar', 'gz', 'jpg', 'jpeg',
        'png', 'gif', 'bmp', 'svg', 'ico', 'mp3', 'mp4', 'avi', 'mov', 'wmv',
        'exe', 'dll', 'so', 'jar', 'war', 'apk', 'ipa', 'deb', 'rpm', 'dmg'
    ]
    
    # 检查是否为文件名（包含文件扩展名）
    if '.' in asset_str:
        parts = asset_str.split('.')
        last_part = parts[-1].lower()
        if last_part in common_file_extensions:
            # 这是一个文件名，不是域名
            return None
    
    domain_pattern = r'^([a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}$'
    if re.match(domain_pattern, asset_str):
        return {
            'asset_type': 'domain',
            'asset_value': asset_str,
            'ip': None,  # 域名不自动填充IP，需要通过DNS解析
            'port': None,
            'protocol': None,
            'url': None
        }
    
    # 匹配域名:端口格式
    domain_port_pattern = r'^([a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}:(\d+)$'
    match = re.match(domain_port_pattern, asset_str)
    if match:
        domain = asset_str.rsplit(':', 1)[0]
        port = int(asset_str.rsplit(':', 1)[1])
        return {
            'asset_type': 'domain',
            'asset_value': asset_str,
            'ip': None,
            'port': port,
            'protocol': None,
            'url': None
        }
    
    return None

@app.route('/api/projects/<int:project_id>/assets', methods=['GET'])
def get_assets(project_id):
    """获取项目的资产列表（支持分页和筛选）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    # 获取分页参数
    page = request.args.get('page', 1, type=int)
    page_size = request.args.get('page_size', 50, type=int)
    
    # 获取筛选参数
    asset_type = request.args.get('asset_type', '')
    asset_value = request.args.get('asset_value', '')
    ip = request.args.get('ip', '')
    port = request.args.get('port', '')
    protocol = request.args.get('protocol', '')
    tag = request.args.get('tag', '')
    notes = request.args.get('notes', '')
    status = request.args.get('status', '')
    risk_level = request.args.get('risk_level', '')
    creator_name = request.args.get('creator_name', '')
    start_date = request.args.get('start_date', '')
    end_date = request.args.get('end_date', '')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该项目
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权访问该项目'}), 403
        
        # 构建WHERE条件
        where_conditions = ['a.project_id = %s']
        params = [project_id]
        
        if asset_type:
            where_conditions.append('a.asset_type = %s')
            params.append(asset_type)
        
        if asset_value:
            where_conditions.append('a.asset_value LIKE %s')
            params.append(f'%{asset_value}%')
        
        if ip:
            where_conditions.append('a.ip LIKE %s')
            params.append(f'%{ip}%')
        
        if port:
            where_conditions.append('a.port = %s')
            params.append(port)
        
        if protocol:
            where_conditions.append('a.protocol = %s')
            params.append(protocol)
        
        if tag:
            where_conditions.append('a.tags LIKE %s')
            params.append(f'%{tag}%')
        
        if notes:
            where_conditions.append('a.notes LIKE %s')
            params.append(f'%{notes}%')
        
        if status:
            where_conditions.append('a.status = %s')
            params.append(status)
        
        if risk_level:
            where_conditions.append('a.risk_level = %s')
            params.append(risk_level)
        
        if creator_name:
            where_conditions.append('u.username LIKE %s')
            params.append(f'%{creator_name}%')
        
        if start_date and end_date:
            where_conditions.append('DATE(a.created_at) BETWEEN %s AND %s')
            params.append(start_date)
            params.append(end_date)
        
        where_clause = ' AND '.join(where_conditions)
        
        # 查询总数
        count_sql = f"""
            SELECT COUNT(*) as total
            FROM assets a
            LEFT JOIN users u ON a.created_by = u.id
            WHERE {where_clause}
        """
        cursor.execute(count_sql, params)
        total = cursor.fetchone()['total']
        
        # 查询分页数据
        offset = (page - 1) * page_size
        data_sql = f"""
            SELECT a.*, u.username as creator_name
            FROM assets a
            LEFT JOIN users u ON a.created_by = u.id
            WHERE {where_clause}
            ORDER BY a.created_at DESC
            LIMIT %s OFFSET %s
        """
        cursor.execute(data_sql, params + [page_size, offset])
        assets = cursor.fetchall()
        cursor.close()
        conn.close()
        
        # 解析tags为数组
        for asset in assets:
            if asset['tags']:
                asset['tags'] = json.loads(asset['tags'])
            else:
                asset['tags'] = []
        
        return jsonify({
            'success': True,
            'assets': assets,
            'total': total,
            'page': page,
            'page_size': page_size,
            'total_pages': (total + page_size - 1) // page_size
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取资产列表失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>/assets', methods=['POST'])
def add_assets(project_id):
    """批量添加资产"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    asset_text = data.get('asset_text', '')
    batch_tags = data.get('batch_tags', [])  # 批量标签
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查项目权限：只有项目所有者或协作者可以添加资产
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此项目'}), 403
        
        # 按行分割资产，然后智能分割每一行
        lines = asset_text.strip().split('\n')
        
        # 提取所有资产项（支持多种分隔符）
        asset_items = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 如果行中包含JSON特征（大括号），则不分割，直接作为一个整体
            if '{' in line or '}' in line:
                asset_items.append(line)
            else:
                # 否则使用正则表达式分割：支持空格、逗号、分号、制表符等
                items = re.split(r'[,;\s\t]+', line)
                for item in items:
                    item = item.strip()
                    if item:
                        asset_items.append(item)
        
        added_count = 0
        skipped_count = 0
        
        for asset_item in asset_items:
            parsed = parse_asset(asset_item)
            if not parsed:
                skipped_count += 1
                continue
            
            # 检查是否已存在
            cursor.execute(
                "SELECT id FROM assets WHERE project_id = %s AND asset_value = %s",
                (project_id, parsed['asset_value'])
            )
            if cursor.fetchone():
                skipped_count += 1
                continue
            
            # 插入资产（默认状态为"未测试"）
            # 将批量标签转换为JSON字符串
            tags_json = json.dumps(batch_tags, ensure_ascii=False) if batch_tags else '[]'
            
            cursor.execute("""
                INSERT INTO assets (project_id, asset_type, asset_value, ip, port, protocol, url, status, tags, created_by)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                project_id,
                parsed['asset_type'],
                parsed['asset_value'],
                parsed['ip'],
                parsed['port'],
                parsed['protocol'],
                parsed['url'],
                '未测试',
                tags_json,
                user_id
            ))
            added_count += 1
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'成功添加 {added_count} 条资产，跳过 {skipped_count} 条',
            'added': added_count,
            'skipped': skipped_count
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'添加资产失败: {str(e)}'}), 500

@app.route('/api/assets/<int:asset_id>', methods=['PUT'])
def update_asset(asset_id):
    """更新资产（标签、备注、状态、风险等级、测试结果）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    tags = data.get('tags', [])
    notes = data.get('notes', '')
    status = data.get('status')
    risk_level = data.get('risk_level')
    test_result = data.get('test_result', '')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：项目所有者或资产创建者可以更新
        if not check_asset_permission(cursor, asset_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此资产'}), 403
        
        cursor.execute(
            "UPDATE assets SET tags = %s, notes = %s, status = %s, risk_level = %s, test_result = %s WHERE id = %s",
            (json.dumps(tags, ensure_ascii=False), notes, status, risk_level, test_result, asset_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '更新成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新失败: {str(e)}'}), 500

@app.route('/api/assets/<int:asset_id>', methods=['DELETE'])
def delete_asset(asset_id):
    """删除资产"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：项目所有者或资产创建者可以删除
        if not check_asset_permission(cursor, asset_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此资产'}), 403
        
        cursor.execute("DELETE FROM assets WHERE id = %s", (asset_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除失败: {str(e)}'}), 500

@app.route('/api/assets/<int:asset_id>/resolve-dns', methods=['POST'])
def resolve_dns(asset_id):
    """解析域名的IP地址"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该资产
        if not check_asset_permission(cursor, asset_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权操作该资产'}), 403
        
        # 获取资产信息
        cursor.execute("SELECT asset_type, asset_value FROM assets WHERE id = %s", (asset_id,))
        asset = cursor.fetchone()
        
        if not asset:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '资产不存在'}), 404
        
        asset_type = asset['asset_type']
        asset_value = asset['asset_value']
        
        # 提取域名
        domain = None
        
        if asset_type == 'domain':
            # 纯域名或域名:端口
            domain = asset_value.split(':')[0] if ':' in asset_value else asset_value
        elif asset_type in ['http', 'https']:
            # 从URL中提取域名
            match = re.match(r'https?://([^:/]+)', asset_value)
            if match:
                domain = match.group(1)
        
        if not domain:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '该资产不是域名类型，无法解析'}), 400
        
        # 检查是否已经是IP地址
        ip_pattern = r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$'
        if re.match(ip_pattern, domain):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '该资产已经是IP地址'}), 400
        
        # DNS解析
        try:
            ip_address = socket.gethostbyname(domain)
            
            # 更新资产的IP字段
            cursor.execute("UPDATE assets SET ip = %s WHERE id = %s", (ip_address, asset_id))
            conn.commit()
            cursor.close()
            conn.close()
            
            return jsonify({
                'success': True,
                'message': f'DNS解析成功',
                'domain': domain,
                'ip': ip_address
            })
        
        except socket.gaierror as e:
            cursor.close()
            conn.close()
            return jsonify({
                'success': False,
                'message': f'DNS解析失败: 无法解析域名 {domain}'
            }), 400
        except Exception as e:
            cursor.close()
            conn.close()
            return jsonify({
                'success': False,
                'message': f'DNS解析失败: {str(e)}'
            }), 500
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'操作失败: {str(e)}'}), 500

@app.route('/api/assets/<int:asset_id>/http-probe', methods=['POST'])
def http_probe(asset_id):
    """HTTP探测：获取网页标题和大小"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该资产
        if not check_asset_permission(cursor, asset_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权操作该资产'}), 403
        
        # 获取资产信息
        cursor.execute("SELECT asset_type, asset_value FROM assets WHERE id = %s", (asset_id,))
        asset = cursor.fetchone()
        
        if not asset:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '资产不存在'}), 404
        
        asset_type = asset['asset_type']
        asset_value = asset['asset_value']
        
        # 构建URL
        url = None
        if asset_type in ['http', 'https']:
            url = asset_value
        elif asset_type == 'ip_port':
            url = f"http://{asset_value}"
        
        if not url:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '该资产类型不支持HTTP探测'}), 400
        
        # 发送HTTP请求
        try:
            response = requests.get(url, timeout=10, verify=False, allow_redirects=True)
            
            # 解析标题
            soup = BeautifulSoup(response.text, 'html.parser')
            title = soup.title.string if soup.title else ''
            content_length = len(response.content)
            
            # 更新资产信息
            cursor.execute(
                "UPDATE assets SET title = %s, content_length = %s WHERE id = %s",
                (title, content_length, asset_id)
            )
            conn.commit()
            cursor.close()
            conn.close()
            
            return jsonify({
                'success': True,
                'message': 'HTTP探测成功',
                'title': title,
                'content_length': content_length,
                'status_code': response.status_code
            })
        
        except requests.exceptions.RequestException as e:
            cursor.close()
            conn.close()
            return jsonify({
                'success': False,
                'message': f'HTTP探测失败: {str(e)}'
            }), 400
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'操作失败: {str(e)}'}), 500

def is_ip_address(s):
    """判断字符串是否为IP地址"""
    ipv4_pattern = r'^(\d{1,3}\.){3}\d{1,3}$'
    ipv6_pattern = r'^([0-9a-fA-F]{0,4}:){7}[0-9a-fA-F]{0,4}$'
    return re.match(ipv4_pattern, s) or re.match(ipv6_pattern, s)

@app.route('/api/projects/<int:project_id>/domain-map', methods=['GET'])
def get_domain_map(project_id):
    """获取项目的域名关系图数据"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, asset_type, asset_value, ip, port, protocol, url
            FROM assets
            WHERE project_id = %s
        """, (project_id,))
        assets = cursor.fetchall()
        cursor.close()
        conn.close()
        
        # 构建节点和边
        nodes = []
        edges = []
        node_map = {}
        
        for asset in assets:
            asset_id = asset['id']
            asset_type = asset['asset_type']
            asset_value = asset['asset_value']
            ip = asset['ip']
            
            # 创建节点
            if asset_type == 'domain':
                domain = asset_value.split(':')[0]
                if domain not in node_map:
                    node_map[domain] = len(nodes)
                    nodes.append({
                        'id': domain,
                        'label': domain,
                        'type': 'domain'
                    })
                
                # 如果有IP，创建IP节点和边
                if ip:
                    if ip not in node_map:
                        node_map[ip] = len(nodes)
                        nodes.append({
                            'id': ip,
                            'label': ip,
                            'type': 'ip'
                        })
                    edges.append({
                        'from': domain,
                        'to': ip,
                        'label': 'resolves to'
                    })
        
        return jsonify({
            'success': True,
            'nodes': nodes,
            'edges': edges
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取域名关系图失败: {str(e)}'}), 500

# ==================== Curl任务管理 ====================

@app.route('/api/projects/<int:project_id>/curl-tasks', methods=['GET'])
def get_curl_tasks(project_id):
    """获取项目的curl任务列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该项目
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权访问该项目'}), 403
        
        cursor.execute("""
            SELECT * FROM curl_tasks
            WHERE project_id = %s
            ORDER BY created_at DESC
        """, (project_id,))
        tasks = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'tasks': tasks})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取任务列表失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>/curl-tasks', methods=['POST'])
def create_curl_task(project_id):
    """创建curl任务"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    data = request.json
    name = data.get('name')
    template_type = data.get('template_type', 'custom')
    template_config = data.get('template_config')
    curl_command = data.get('curl_command')
    extract_pattern = data.get('extract_pattern')
    batch_tags = data.get('batch_tags', [])
    schedule_type = data.get('schedule_type', 'manual')
    schedule_config = data.get('schedule_config')
    enabled = data.get('enabled', True)
    
    if not name:
        return jsonify({'success': False, 'message': '任务名称不能为空'}), 400
    
    if template_type == 'custom' and not curl_command:
        return jsonify({'success': False, 'message': '自定义模板需要提供curl命令'}), 400
    
    if template_type == 'lighthouse' and not template_config:
        return jsonify({'success': False, 'message': '灯塔导入模板需要提供配置信息'}), 400
    
    if template_type == 'fofa' and not template_config:
        return jsonify({'success': False, 'message': 'FOFA查询模板需要提供配置信息'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该项目
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权在该项目中创建任务'}), 403
        
        schedule_config_json = json.dumps(schedule_config) if schedule_config else None
        template_config_json = json.dumps(template_config) if template_config else None
        batch_tags_json = json.dumps(batch_tags, ensure_ascii=False) if batch_tags else None
        
        cursor.execute("""
            INSERT INTO curl_tasks 
            (project_id, name, template_type, template_config, curl_command, extract_pattern, batch_tags, schedule_type, schedule_config, enabled, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (project_id, name, template_type, template_config_json, curl_command, extract_pattern, batch_tags_json, schedule_type, schedule_config_json, enabled, user_id))
        
        conn.commit()
        task_id = cursor.lastrowid
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '任务创建成功',
            'task_id': task_id
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建任务失败: {str(e)}'}), 500

# ==================== 大屏数据接口 ====================

@app.route('/api/dashboard/bigscreen', methods=['GET'])
def get_bigscreen_data():
    """获取大屏展示数据"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取当前用户及其协作人的ID列表
        cursor.execute("""
            SELECT collaborator_id FROM user_collaborations WHERE user_id = %s
        """, (user_id,))
        collaborator_rows = cursor.fetchall()
        collaborator_ids = [row['collaborator_id'] for row in collaborator_rows]
        
        # 包含当前用户自己
        all_user_ids = [user_id] + collaborator_ids
        
        # 构建IN查询的占位符
        placeholders = ','.join(['%s'] * len(all_user_ids))
        
        # 总资产数（只统计当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT COUNT(*) as count FROM assets 
            WHERE project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
        """, all_user_ids)
        total_assets = cursor.fetchone()['count']
        
        # 总项目数（只统计当前用户及协作人创建的项目）
        cursor.execute(f"""
            SELECT COUNT(*) as count FROM projects 
            WHERE created_by IN ({placeholders})
        """, all_user_ids)
        total_projects = cursor.fetchone()['count']
        
        # 用户数量（当前用户 + 协作人数量）
        total_users = len(all_user_ids)
        
        # 今日新增资产（只统计当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT COUNT(*) as count FROM assets 
            WHERE DATE(created_at) = CURDATE()
            AND project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
        """, all_user_ids)
        today_assets = cursor.fetchone()['count']
        
        # 资产类型分布（只统计当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT asset_type as type, COUNT(*) as count 
            FROM assets 
            WHERE project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
            GROUP BY asset_type
        """, all_user_ids)
        asset_type_data = cursor.fetchall()
        
        # 转换资产类型为中文
        type_map = {
            'ip': 'IP地址',
            'ip_port': 'IP:端口',
            'domain': '域名',
            'http': 'HTTP',
            'https': 'HTTPS'
        }
        asset_type_distribution = [
            {'type': type_map.get(item['type'], item['type']), 'count': item['count']}
            for item in asset_type_data
        ]
        
        # 风险等级分布（只统计当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT 
                COALESCE(risk_level, '信息') as level,
                COUNT(*) as count 
            FROM assets 
            WHERE project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
            GROUP BY risk_level
        """, all_user_ids)
        risk_level_distribution = cursor.fetchall()
        
        # 最近7天资产趋势（只统计当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT 
                DATE_FORMAT(DATE(created_at), '%%m-%%d') as date,
                COUNT(*) as count
            FROM assets
            WHERE created_at >= DATE_SUB(CURDATE(), INTERVAL 7 DAY)
            AND project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
            GROUP BY DATE_FORMAT(DATE(created_at), '%%m-%%d'), DATE(created_at)
            ORDER BY DATE(created_at)
        """, all_user_ids)
        asset_trend = cursor.fetchall()
        
        # 实时动态（最近10条资产，只显示当前用户及协作人的项目下的资产）
        cursor.execute(f"""
            SELECT 
                a.asset_value,
                a.asset_type,
                a.created_at,
                p.name as project_name
            FROM assets a
            LEFT JOIN projects p ON a.project_id = p.id
            WHERE p.created_by IN ({placeholders})
            ORDER BY a.created_at DESC
            LIMIT 10
        """, all_user_ids)
        recent_assets = cursor.fetchall()
        
        recent_activities = []
        for asset in recent_assets:
            time_diff = datetime.now() - asset['created_at']
            if time_diff.seconds < 60:
                time_str = f"{time_diff.seconds}秒前"
            elif time_diff.seconds < 3600:
                time_str = f"{time_diff.seconds // 60}分钟前"
            else:
                time_str = asset['created_at'].strftime('%H:%M')
            
            recent_activities.append({
                'text': f"新增{asset['asset_type']}资产: {asset['asset_value'][:30]}",
                'time': time_str
            })
        
        # Top5项目（按资产数量，只统计当前用户及协作人的项目）
        cursor.execute(f"""
            SELECT 
                p.id,
                p.name,
                COUNT(a.id) as asset_count
            FROM projects p
            LEFT JOIN assets a ON p.id = a.project_id
            WHERE p.created_by IN ({placeholders})
            GROUP BY p.id, p.name
            ORDER BY asset_count DESC
            LIMIT 5
        """, all_user_ids)
        top_projects = cursor.fetchall()
        
        # 漏洞统计数据（只统计当前用户及协作人的项目下的漏洞）
        # 总漏洞数
        cursor.execute(f"""
            SELECT COUNT(*) as count FROM vulnerabilities 
            WHERE project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
        """, all_user_ids)
        total_vulnerabilities = cursor.fetchone()['count']
        
        # 各状态漏洞数量
        cursor.execute(f"""
            SELECT status, COUNT(*) as count 
            FROM vulnerabilities 
            WHERE project_id IN (
                SELECT id FROM projects WHERE created_by IN ({placeholders})
            )
            GROUP BY status
        """, all_user_ids)
        vuln_status_data = cursor.fetchall()
        
        # 初始化各状态数量
        vuln_stats = {
            '已提交': 0,
            '未提交': 0,
            '已修复': 0,
            '未修复': 0,
            '无法访问': 0,
            '无法复现': 0
        }
        
        # 填充实际数据
        for item in vuln_status_data:
            if item['status'] in vuln_stats:
                vuln_stats[item['status']] = item['count']
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'data': {
                'totalAssets': total_assets,
                'totalProjects': total_projects,
                'totalUsers': total_users,
                'todayAssets': today_assets,
                'assetTypeDistribution': asset_type_distribution,
                'riskLevelDistribution': risk_level_distribution,
                'assetTrend': asset_trend,
                'recentActivities': recent_activities,
                'topProjects': top_projects,
                'totalVulnerabilities': total_vulnerabilities,
                'vulnSubmitted': vuln_stats['已提交'],
                'vulnNotSubmitted': vuln_stats['未提交'],
                'vulnFixed': vuln_stats['已修复'],
                'vulnNotFixed': vuln_stats['未修复'],
                'vulnNoAccess': vuln_stats['无法访问'],
                'vulnCannotReproduce': vuln_stats['无法复现']
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取大屏数据失败: {str(e)}'}), 500

# ==================== 共享凭证管理接口 ====================

@app.route('/api/credentials', methods=['GET'])
def get_credentials():
    """获取共享凭证列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    # 获取筛选参数
    category = request.args.get('category', '')
    platform = request.args.get('platform', '')
    is_active = request.args.get('is_active', '')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 构建WHERE条件
        where_conditions = []
        params = []
        
        if category:
            where_conditions.append('c.category = %s')
            params.append(category)
        
        if platform:
            where_conditions.append('c.platform LIKE %s')
            params.append(f'%{platform}%')
        
        if is_active:
            where_conditions.append('c.is_active = %s')
            params.append(is_active == 'true')
        
        where_clause = ' AND '.join(where_conditions) if where_conditions else '1=1'
        
        cursor.execute(f"""
            SELECT c.*, u.username as creator_name
            FROM shared_credentials c
            LEFT JOIN users u ON c.created_by = u.id
            WHERE {where_clause}
            ORDER BY c.created_at DESC
        """, params)
        credentials = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'credentials': credentials})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取凭证列表失败: {str(e)}'}), 500

@app.route('/api/credentials', methods=['POST'])
def create_credential():
    """创建共享凭证"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    data = request.json
    platform = data.get('platform')
    url = data.get('url', '')
    username = data.get('username')
    password = data.get('password')
    category = data.get('category', '其他')
    description = data.get('description', '')
    is_active = data.get('is_active', True)
    
    if not platform or not username or not password:
        return jsonify({'success': False, 'message': '平台名称、用户名和密码不能为空'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 简单的密码加密（实际生产环境应使用更安全的加密方式）
        import base64
        encrypted_password = base64.b64encode(password.encode()).decode()
        
        cursor.execute("""
            INSERT INTO shared_credentials 
            (platform, url, username, password, category, description, is_active, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (platform, url, username, encrypted_password, category, description, is_active, user_id))
        
        credential_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '凭证创建成功', 'credential_id': credential_id})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建凭证失败: {str(e)}'}), 500

@app.route('/api/credentials/<int:credential_id>', methods=['PUT'])
def update_credential(credential_id):
    """更新共享凭证"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：只有凭证创建者才能更新
        cursor.execute("SELECT created_by FROM shared_credentials WHERE id = %s", (credential_id,))
        credential = cursor.fetchone()
        if not credential:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '凭证不存在'}), 404
        
        if credential['created_by'] != user_id:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此凭证'}), 403
        
        # 构建更新语句
        update_fields = []
        update_values = []
        
        if 'platform' in data:
            update_fields.append('platform = %s')
            update_values.append(data['platform'])
        
        if 'url' in data:
            update_fields.append('url = %s')
            update_values.append(data['url'])
        
        if 'username' in data:
            update_fields.append('username = %s')
            update_values.append(data['username'])
        
        if 'password' in data:
            import base64
            encrypted_password = base64.b64encode(data['password'].encode()).decode()
            update_fields.append('password = %s')
            update_values.append(encrypted_password)
        
        if 'category' in data:
            update_fields.append('category = %s')
            update_values.append(data['category'])
        
        if 'description' in data:
            update_fields.append('description = %s')
            update_values.append(data['description'])
        
        if 'is_active' in data:
            update_fields.append('is_active = %s')
            update_values.append(data['is_active'])
        
        if not update_fields:
            return jsonify({'success': False, 'message': '没有要更新的字段'}), 400
        
        update_values.append(credential_id)
        
        cursor.execute(f"""
            UPDATE shared_credentials 
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, update_values)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '凭证更新成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新凭证失败: {str(e)}'}), 500

@app.route('/api/credentials/<int:credential_id>', methods=['DELETE'])
def delete_credential(credential_id):
    """删除共享凭证"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：只有凭证创建者才能删除
        cursor.execute("SELECT created_by FROM shared_credentials WHERE id = %s", (credential_id,))
        credential = cursor.fetchone()
        if not credential:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '凭证不存在'}), 404
        
        if credential['created_by'] != user_id:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此凭证'}), 403
        
        cursor.execute("DELETE FROM shared_credentials WHERE id = %s", (credential_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '凭证删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除凭证失败: {str(e)}'}), 500

@app.route('/api/credentials/<int:credential_id>/password', methods=['GET'])
def get_credential_password(credential_id):
    """获取凭证密码（解密）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT password, created_by FROM shared_credentials WHERE id = %s", (credential_id,))
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not result:
            return jsonify({'success': False, 'message': '凭证不存在'}), 404
        
        # 注意：共享凭证对所有用户可见，但只有创建者可以查看密码
        # 如果需要所有用户都能查看，可以去掉这个检查
        if result['created_by'] != user_id:
            return jsonify({'success': False, 'message': '无权限查看此凭证密码'}), 403
        
        # 解密密码
        import base64
        decrypted_password = base64.b64decode(result['password'].encode()).decode()
        
        return jsonify({'success': True, 'password': decrypted_password})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取密码失败: {str(e)}'}), 500

# ==================== 漏洞生命周期管理接口 ====================

@app.route('/api/projects/<int:project_id>/vulnerabilities', methods=['GET'])
def get_vulnerabilities(project_id):
    """获取项目的漏洞列表（支持筛选）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    # 获取筛选参数
    vuln_type = request.args.get('vuln_type', '')
    status = request.args.get('status', '')
    discoverer = request.args.get('discoverer', '')
    url_filter = request.args.get('url', '')
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该项目
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权访问该项目'}), 403
        
        # 构建WHERE条件
        where_conditions = ['v.project_id = %s']
        params = [project_id]
        
        if vuln_type:
            where_conditions.append('v.vuln_type LIKE %s')
            params.append(f'%{vuln_type}%')
        
        if status:
            where_conditions.append('v.status = %s')
            params.append(status)
        
        if discoverer:
            where_conditions.append('v.discoverer LIKE %s')
            params.append(f'%{discoverer}%')
        
        if url_filter:
            where_conditions.append('v.url LIKE %s')
            params.append(f'%{url_filter}%')
        
        where_clause = ' AND '.join(where_conditions)
        
        cursor.execute(f"""
            SELECT v.*, u.username as creator_name
            FROM vulnerabilities v
            LEFT JOIN users u ON v.created_by = u.id
            WHERE {where_clause}
            ORDER BY v.created_at DESC
        """, params)
        vulnerabilities = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'vulnerabilities': vulnerabilities})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取漏洞列表失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>/vulnerabilities', methods=['POST'])
def create_vulnerability(project_id):
    """创建漏洞"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    url = data.get('url')
    vuln_type = data.get('vuln_type')
    description = data.get('description', '')
    discoverer = data.get('discoverer', '')
    markdown_detail = data.get('markdown_detail', '')
    status = data.get('status', '未提交')
    
    if not url or not vuln_type:
        return jsonify({'success': False, 'message': 'URL和漏洞类型不能为空'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 权限检查：验证用户是否有权访问该项目
        if not check_project_permission(cursor, project_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权在该项目中创建漏洞'}), 403
        
        cursor.execute("""
            INSERT INTO vulnerabilities 
            (project_id, url, vuln_type, description, discoverer, markdown_detail, status, created_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (project_id, url, vuln_type, description, discoverer, markdown_detail, status, user_id))
        
        vuln_id = cursor.lastrowid
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '漏洞创建成功', 'vulnerability_id': vuln_id})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建漏洞失败: {str(e)}'}), 500

@app.route('/api/vulnerabilities/<int:vuln_id>', methods=['PUT'])
def update_vulnerability(vuln_id):
    """更新漏洞"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：项目所有者或漏洞创建者可以更新
        if not check_vulnerability_permission(cursor, vuln_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此漏洞'}), 403
        
        # 构建更新语句
        update_fields = []
        update_values = []
        
        if 'url' in data:
            update_fields.append('url = %s')
            update_values.append(data['url'])
        
        if 'vuln_type' in data:
            update_fields.append('vuln_type = %s')
            update_values.append(data['vuln_type'])
        
        if 'description' in data:
            update_fields.append('description = %s')
            update_values.append(data['description'])
        
        if 'discoverer' in data:
            update_fields.append('discoverer = %s')
            update_values.append(data['discoverer'])
        
        if 'markdown_detail' in data:
            update_fields.append('markdown_detail = %s')
            update_values.append(data['markdown_detail'])
        
        if 'status' in data:
            update_fields.append('status = %s')
            update_values.append(data['status'])
        
        if not update_fields:
            return jsonify({'success': False, 'message': '没有要更新的字段'}), 400
        
        update_values.append(vuln_id)
        
        cursor.execute(f"""
            UPDATE vulnerabilities 
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, update_values)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '漏洞更新成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新漏洞失败: {str(e)}'}), 500

@app.route('/api/vulnerabilities/<int:vuln_id>', methods=['DELETE'])
def delete_vulnerability(vuln_id):
    """删除漏洞"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查权限：项目所有者或漏洞创建者可以删除
        if not check_vulnerability_permission(cursor, vuln_id, user_id, require_owner=False):
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作此漏洞'}), 403
        
        cursor.execute("DELETE FROM vulnerabilities WHERE id = %s", (vuln_id,))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '漏洞删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除漏洞失败: {str(e)}'}), 500

# ==================== 超级管理员功能 ====================

def check_super_admin(user_id):
    """检查用户是否为超级管理员"""
    conn = get_db_connection()
    if not conn:
        return False
    
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT is_super_admin FROM users WHERE id = %s",
            (user_id,)
        )
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        return result and result.get('is_super_admin', False)
    except:
        return False

@app.route('/api/admin/users', methods=['GET'])
def admin_get_users():
    """超级管理员：获取所有用户列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限：必须验证当前登录用户，不能从请求参数获取
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取所有用户及其统计信息
        cursor.execute("""
            SELECT 
                u.id,
                u.username,
                u.is_admin,
                u.is_super_admin,
                u.created_at,
                COUNT(DISTINCT p.id) as project_count,
                COUNT(DISTINCT a.id) as asset_count
            FROM users u
            LEFT JOIN projects p ON u.id = p.created_by
            LEFT JOIN assets a ON u.id = a.created_by
            GROUP BY u.id, u.username, u.is_admin, u.is_super_admin, u.created_at
            ORDER BY u.created_at DESC
        """)
        
        users = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'users': users
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取用户列表失败: {str(e)}'}), 500

@app.route('/api/admin/users', methods=['POST'])
def admin_create_user():
    """超级管理员：创建新用户"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '无权限访问'}), 403
    
    data = request.json
    username = data.get('username')
    password = data.get('password')
    is_super_admin = data.get('is_super_admin', False)
    
    # 验证必填字段
    if not username or not password:
        return jsonify({'success': False, 'message': '用户名和密码不能为空'}), 400
    
    if len(username) < 3 or len(username) > 50:
        return jsonify({'success': False, 'message': '用户名长度必须在3-50个字符之间'}), 400
    
    if len(password) < 6:
        return jsonify({'success': False, 'message': '密码长度至少为6个字符'}), 400
    
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查用户名是否已存在
        cursor.execute("SELECT id FROM users WHERE username = %s", (username,))
        if cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '用户名已存在'}), 400
        
        # 创建用户
        cursor.execute("""
            INSERT INTO users (username, password, is_super_admin) 
            VALUES (%s, %s, %s)
        """, (username, password_hash, is_super_admin))
        
        user_id = cursor.lastrowid
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True, 
            'message': '用户创建成功',
            'user': {
                'id': user_id,
                'username': username,
                'is_super_admin': is_super_admin
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'创建用户失败: {str(e)}'}), 500

@app.route('/api/admin/users/<int:target_user_id>', methods=['PUT'])
def admin_update_user(target_user_id):
    """超级管理员：更新用户信息（设置超级管理员权限、修改用户名、密码等）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '无权限访问'}), 403
    
    data = request.json
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        update_fields = []
        update_values = []
        
        # 可更新的字段：超级管理员权限
        if 'is_super_admin' in data:
            # 防止最后一个超级管理员被移除
            if not data['is_super_admin']:
                cursor.execute("SELECT COUNT(*) as count FROM users WHERE is_super_admin = TRUE")
                result = cursor.fetchone()
                if result['count'] <= 1:
                    cursor.close()
                    conn.close()
                    return jsonify({'success': False, 'message': '至少需要保留一个超级管理员'}), 400
            
            update_fields.append('is_super_admin = %s')
            update_values.append(data['is_super_admin'])
        
        # 可更新的字段：用户名
        if 'username' in data:
            new_username = data['username']
            if len(new_username) < 3 or len(new_username) > 50:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '用户名长度必须在3-50个字符之间'}), 400
            
            # 检查用户名是否已被其他用户使用
            cursor.execute("SELECT id FROM users WHERE username = %s AND id != %s", (new_username, target_user_id))
            if cursor.fetchone():
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '用户名已被使用'}), 400
            
            update_fields.append('username = %s')
            update_values.append(new_username)
        
        # 可更新的字段：密码
        if 'password' in data:
            new_password = data['password']
            if len(new_password) < 6:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '密码长度至少为6个字符'}), 400
            
            password_hash = hashlib.sha256(new_password.encode()).hexdigest()
            update_fields.append('password = %s')
            update_values.append(password_hash)
        
        if not update_fields:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '没有要更新的字段'}), 400
        
        update_values.append(target_user_id)
        
        cursor.execute(f"""
            UPDATE users 
            SET {', '.join(update_fields)}
            WHERE id = %s
        """, update_values)
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '用户信息更新成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新用户信息失败: {str(e)}'}), 500

@app.route('/api/admin/users/<int:target_user_id>', methods=['DELETE'])
def admin_delete_user(target_user_id):
    """超级管理员：删除用户"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '无权限访问'}), 403
    
    # 不能删除自己
    if user['id'] == target_user_id:
        return jsonify({'success': False, 'message': '不能删除自己的账户'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查是否为超级管理员
        cursor.execute("SELECT is_super_admin FROM users WHERE id = %s", (target_user_id,))
        user = cursor.fetchone()
        
        if user and user['is_super_admin']:
            # 防止删除最后一个超级管理员
            cursor.execute("SELECT COUNT(*) as count FROM users WHERE is_super_admin = TRUE")
            result = cursor.fetchone()
            if result['count'] <= 1:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '不能删除最后一个超级管理员'}), 400
        
        # 删除用户（级联删除相关数据）
        cursor.execute("DELETE FROM users WHERE id = %s", (target_user_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '用户删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除用户失败: {str(e)}'}), 500

@app.route('/api/admin/projects', methods=['GET'])
def admin_get_all_projects():
    """超级管理员：获取所有项目列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '无权限访问'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取所有项目及其统计信息
        cursor.execute("""
            SELECT 
                p.id,
                p.name,
                p.description,
                p.created_by,
                u.username as creator_name,
                p.created_at,
                p.updated_at,
                COUNT(DISTINCT a.id) as asset_count
            FROM projects p
            LEFT JOIN users u ON p.created_by = u.id
            LEFT JOIN assets a ON p.id = a.project_id
            GROUP BY p.id, p.name, p.description, p.created_by, u.username, p.created_at, p.updated_at
            ORDER BY p.created_at DESC
        """)
        
        projects = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'projects': projects
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取项目列表失败: {str(e)}'}), 500

@app.route('/api/admin/projects/<int:project_id>', methods=['DELETE'])
def admin_delete_project(project_id):
    """超级管理员：删除任意项目"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限：必须验证当前登录用户，不能从请求参数获取
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 删除项目（级联删除相关资产）
        cursor.execute("DELETE FROM projects WHERE id = %s", (project_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '项目删除成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除项目失败: {str(e)}'}), 500

@app.route('/api/admin/settings', methods=['GET'])
def admin_get_settings():
    """超级管理员：获取系统设置"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限：必须验证当前登录用户，不能从请求参数获取
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取所有系统设置
        cursor.execute("SELECT * FROM system_settings ORDER BY setting_key")
        settings = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # 转换为字典格式
        settings_dict = {}
        for setting in settings:
            settings_dict[setting['setting_key']] = {
                'value': setting['setting_value'],
                'description': setting['description']
            }
        
        return jsonify({
            'success': True,
            'settings': settings_dict
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取系统设置失败: {str(e)}'}), 500

@app.route('/api/admin/settings', methods=['PUT'])
def admin_update_settings():
    """超级管理员：更新系统设置"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限：必须验证当前登录用户，不能从请求参数获取
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
    
    data = request.json
    settings = data.get('settings', {})
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 更新每个设置
        for key, value in settings.items():
            cursor.execute("""
                UPDATE system_settings 
                SET setting_value = %s 
                WHERE setting_key = %s
            """, (value, key))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '系统设置更新成功'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新系统设置失败: {str(e)}'}), 500

@app.route('/api/settings/registration', methods=['GET'])
def get_registration_setting():
    """获取注册开关状态（公开接口）"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT setting_value 
            FROM system_settings 
            WHERE setting_key = 'allow_registration'
        """)
        
        result = cursor.fetchone()
        cursor.close()
        conn.close()
        
        # 默认允许注册
        allow_registration = True
        if result:
            allow_registration = result['setting_value'].lower() == 'true'
        
        return jsonify({
            'success': True,
            'allow_registration': allow_registration
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取设置失败: {str(e)}'}), 500

@app.route('/api/admin/stats', methods=['GET'])
def admin_get_stats():
    """超级管理员：获取系统统计信息"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    # 验证超级管理员权限：必须验证当前登录用户，不能从请求参数获取
    if not user.get('is_super_admin', False):
        return jsonify({'success': False, 'message': '需要超级管理员权限'}), 403
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取用户总数
        cursor.execute("SELECT COUNT(*) as count FROM users")
        user_count = cursor.fetchone()['count']
        
        # 获取项目总数
        cursor.execute("SELECT COUNT(*) as count FROM projects")
        project_count = cursor.fetchone()['count']
        
        # 获取资产总数
        cursor.execute("SELECT COUNT(*) as count FROM assets")
        asset_count = cursor.fetchone()['count']
        
        # 获取漏洞总数
        cursor.execute("SELECT COUNT(*) as count FROM vulnerabilities")
        vuln_count = cursor.fetchone()['count']
        
        # 获取最近7天新增用户数
        cursor.execute("""
            SELECT COUNT(*) as count 
            FROM users 
            WHERE created_at >= DATE_SUB(NOW(), INTERVAL 7 DAY)
        """)
        recent_users = cursor.fetchone()['count']
        
        # 获取最近7天新增项目数
        cursor.execute("""
            SELECT COUNT(*) as count 
            FROM projects 
            WHERE created_at >= DATE_SUB(NOW(), INTERVAL 7 DAY)
        """)
        recent_projects = cursor.fetchone()['count']
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'stats': {
                'total_users': user_count,
                'total_projects': project_count,
                'total_assets': asset_count,
                'total_vulnerabilities': vuln_count,
                'recent_users': recent_users,
                'recent_projects': recent_projects
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取统计信息失败: {str(e)}'}), 500

# ==================== 协同任务分配 ====================

@app.route('/api/users/<int:user_id>/invitees', methods=['GET'])
def get_invitees(user_id):
    """获取通过邀请码注册的用户列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 从 user_collaborations 表获取协作者（被邀请人）
        cursor.execute("""
            SELECT u.id, u.username, uc.created_at
            FROM user_collaborations uc
            JOIN users u ON uc.collaborator_id = u.id
            WHERE uc.user_id = %s
            ORDER BY uc.created_at DESC
        """, (user_id,))
        invitees = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'invitees': invitees
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取被邀请人列表失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>/visibility', methods=['PUT'])
def update_project_visibility(project_id):
    """更新项目对邀请人的可见性"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    is_visible = data.get('is_visible_to_inviter', True)
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 验证项目所有权
        cursor.execute("SELECT created_by FROM projects WHERE id = %s", (project_id,))
        project = cursor.fetchone()
        
        if not project:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '项目不存在'}), 404
        
        if project['created_by'] != user_id:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限修改此项目'}), 403
        
        # 更新可见性
        cursor.execute("""
            UPDATE projects 
            SET is_visible_to_inviter = %s 
            WHERE id = %s
        """, (is_visible, project_id))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '项目可见性更新成功'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'更新项目可见性失败: {str(e)}'}), 500

@app.route('/api/projects/<int:project_id>/assign-task', methods=['POST'])
def assign_collaboration_task(project_id):
    """下发协同任务"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    assignee_ids = data.get('assignee_ids', [])  # 被分配人ID列表
    asset_ids = data.get('asset_ids', [])  # 下发的资产ID列表
    task_description = data.get('task_description', '')
    source_project_visible = data.get('source_project_visible', True)  # 源项目是否对接收人可见
    
    if not assignee_ids:
        return jsonify({'success': False, 'message': '请选择任务接收人'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 检查协同任务相关字段是否存在
        cursor.execute("SHOW COLUMNS FROM projects LIKE 'is_collaboration_task'")
        has_collaboration_fields = cursor.fetchone() is not None
        
        if not has_collaboration_fields:
            cursor.close()
            conn.close()
            return jsonify({
                'success': False, 
                'message': '协同任务功能未启用，请先运行数据库迁移: python backend/migrate_add_collaboration_tasks.py'
            }), 400
        
        # 验证项目所有权
        cursor.execute("SELECT created_by, name FROM projects WHERE id = %s", (project_id,))
        project = cursor.fetchone()
        
        if not project:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '项目不存在'}), 404
        
        if project['created_by'] != user_id:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限下发此项目的任务'}), 403
        
        created_tasks = []
        
        for assignee_id in assignee_ids:
            # 验证被分配人是否是协作者
            cursor.execute("""
                SELECT COUNT(*) as count 
                FROM user_collaborations 
                WHERE user_id = %s AND collaborator_id = %s
            """, (user_id, assignee_id))
            
            if cursor.fetchone()['count'] == 0:
                continue  # 跳过非协作者
            
            # 为被分配人创建协同任务项目
            cursor.execute("""
                INSERT INTO projects (name, description, created_by, is_collaboration_task, source_project_id, task_assigner_id)
                VALUES (%s, %s, %s, TRUE, %s, %s)
            """, (
                f"[协同任务] {project['name']}", 
                task_description or f"来自 {user['username']} 的协同任务",
                assignee_id,
                project_id,
                user_id
            ))
            
            target_project_id = cursor.lastrowid
            
            # 创建协同任务记录
            cursor.execute("""
                INSERT INTO collaboration_tasks 
                (source_project_id, target_project_id, assigner_id, assignee_id, asset_ids, task_description, status)
                VALUES (%s, %s, %s, %s, %s, %s, 'active')
            """, (
                project_id,
                target_project_id,
                user_id,
                assignee_id,
                json.dumps(asset_ids) if asset_ids else None,
                task_description
            ))
            
            collaboration_task_id = cursor.lastrowid
            
            # 设置源项目对接收人的可见性
            if not source_project_visible:
                # 在权限表中记录：该用户不能查看源项目
                cursor.execute("""
                    INSERT INTO project_access_permissions (project_id, user_id, can_view)
                    VALUES (%s, %s, FALSE)
                    ON DUPLICATE KEY UPDATE can_view = FALSE
                """, (project_id, assignee_id))
            
            # 检查 source_asset_id 字段是否存在
            cursor.execute("SHOW COLUMNS FROM assets LIKE 'source_asset_id'")
            has_source_asset_field = cursor.fetchone() is not None
            
            # 如果指定了资产，复制资产到目标项目
            if asset_ids:
                for asset_id in asset_ids:
                    # 获取源资产信息
                    cursor.execute("SELECT * FROM assets WHERE id = %s AND project_id = %s", (asset_id, project_id))
                    source_asset = cursor.fetchone()
                    
                    if source_asset:
                        # 复制资产到目标项目
                        if has_source_asset_field:
                            cursor.execute("""
                                INSERT INTO assets 
                                (project_id, asset_type, asset_value, ip, port, protocol, url, title, content_length, 
                                 status, risk_level, tags, notes, test_result, created_by, source_asset_id)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                                target_project_id,
                                source_asset['asset_type'],
                                source_asset['asset_value'],
                                source_asset['ip'],
                                source_asset['port'],
                                source_asset['protocol'],
                                source_asset['url'],
                                source_asset['title'],
                                source_asset['content_length'],
                                source_asset['status'],
                                source_asset['risk_level'],
                                source_asset['tags'],
                                source_asset['notes'],
                                source_asset['test_result'],
                                assignee_id,
                                asset_id
                            ))
                        else:
                            cursor.execute("""
                                INSERT INTO assets 
                                (project_id, asset_type, asset_value, ip, port, protocol, url, title, content_length, 
                                 status, risk_level, tags, notes, test_result, created_by)
                                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                                target_project_id,
                                source_asset['asset_type'],
                                source_asset['asset_value'],
                                source_asset['ip'],
                                source_asset['port'],
                                source_asset['protocol'],
                                source_asset['url'],
                                source_asset['title'],
                                source_asset['content_length'],
                                source_asset['status'],
                                source_asset['risk_level'],
                                source_asset['tags'],
                                source_asset['notes'],
                                source_asset['test_result'],
                                assignee_id
                            ))
            
            # 获取被分配人用户名
            cursor.execute("SELECT username FROM users WHERE id = %s", (assignee_id,))
            assignee = cursor.fetchone()
            
            created_tasks.append({
                'assignee_id': assignee_id,
                'assignee_name': assignee['username'] if assignee else 'Unknown',
                'target_project_id': target_project_id,
                'collaboration_task_id': collaboration_task_id
            })
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'成功下发任务给 {len(created_tasks)} 个协作者',
            'tasks': created_tasks
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'下发任务失败: {str(e)}'}), 500

@app.route('/api/collaboration-tasks/<int:task_id>/sync', methods=['POST'])
def sync_collaboration_task(task_id):
    """同步协同任务的资产变更"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    sync_type = data.get('sync_type')  # 'asset_add', 'asset_update', 'asset_delete'
    asset_id = data.get('asset_id')
    sync_data = data.get('sync_data', {})
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 获取协同任务信息
        cursor.execute("""
            SELECT source_project_id, target_project_id, assigner_id, assignee_id
            FROM collaboration_tasks
            WHERE id = %s AND status = 'active'
        """, (task_id,))
        
        task = cursor.fetchone()
        if not task:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '协同任务不存在或已结束'}), 404
        
        # 确定同步方向
        if user_id == task['assignee_id']:
            sync_direction = 'to_source'
            source_project = task['target_project_id']
            target_project = task['source_project_id']
        elif user_id == task['assigner_id']:
            sync_direction = 'to_target'
            source_project = task['source_project_id']
            target_project = task['target_project_id']
        else:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限同步此任务'}), 403
        
        # 记录同步日志
        cursor.execute("""
            INSERT INTO project_sync_logs 
            (collaboration_task_id, sync_type, asset_id, sync_data, sync_direction)
            VALUES (%s, %s, %s, %s, %s)
        """, (task_id, sync_type, asset_id, json.dumps(sync_data), sync_direction))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '同步记录已创建'
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'同步失败: {str(e)}'}), 500

@app.route('/api/collaboration-tasks', methods=['GET'])
def get_collaboration_tasks():
    """获取协同任务列表"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    task_type = request.args.get('type', 'all')  # 'assigned', 'received', 'all'
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        if task_type == 'assigned':
            # 我下发的任务
            cursor.execute("""
                SELECT ct.*, 
                       sp.name as source_project_name,
                       tp.name as target_project_name,
                       u.username as assignee_name
                FROM collaboration_tasks ct
                LEFT JOIN projects sp ON ct.source_project_id = sp.id
                LEFT JOIN projects tp ON ct.target_project_id = tp.id
                LEFT JOIN users u ON ct.assignee_id = u.id
                WHERE ct.assigner_id = %s
                ORDER BY ct.created_at DESC
            """, (user_id,))
        elif task_type == 'received':
            # 我接收的任务
            cursor.execute("""
                SELECT ct.*, 
                       sp.name as source_project_name,
                       tp.name as target_project_name,
                       u.username as assigner_name
                FROM collaboration_tasks ct
                LEFT JOIN projects sp ON ct.source_project_id = sp.id
                LEFT JOIN projects tp ON ct.target_project_id = tp.id
                LEFT JOIN users u ON ct.assigner_id = u.id
                WHERE ct.assignee_id = %s
                ORDER BY ct.created_at DESC
            """, (user_id,))
        else:
            # 所有相关任务
            cursor.execute("""
                SELECT ct.*, 
                       sp.name as source_project_name,
                       tp.name as target_project_name,
                       assigner.username as assigner_name,
                       assignee.username as assignee_name
                FROM collaboration_tasks ct
                LEFT JOIN projects sp ON ct.source_project_id = sp.id
                LEFT JOIN projects tp ON ct.target_project_id = tp.id
                LEFT JOIN users assigner ON ct.assigner_id = assigner.id
                LEFT JOIN users assignee ON ct.assignee_id = assignee.id
                WHERE ct.assigner_id = %s OR ct.assignee_id = %s
                ORDER BY ct.created_at DESC
            """, (user_id, user_id))
        
        tasks = cursor.fetchall()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'tasks': tasks
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'获取协同任务列表失败: {str(e)}'}), 500

@app.route('/api/projects/merge', methods=['POST'])
def merge_projects():
    """合并多个项目到一个项目"""
    if not check_installed():
        return jsonify({'success': False, 'message': '系统未安装'}), 400
    
    # 获取当前用户（安全方式）
    user = get_current_user()
    if not user:
        return jsonify({'success': False, 'message': '未登录或认证已过期'}), 401
    
    user_id = user['id']
    
    data = request.json
    source_project_ids = data.get('source_project_ids', [])  # 源项目ID列表
    target_project_id = data.get('target_project_id')  # 目标项目ID
    delete_source = data.get('delete_source', False)  # 是否删除源项目
    
    if not source_project_ids or len(source_project_ids) == 0:
        return jsonify({'success': False, 'message': '请选择要合并的源项目'}), 400
    
    if not target_project_id:
        return jsonify({'success': False, 'message': '请选择目标项目'}), 400
    
    if target_project_id in source_project_ids:
        return jsonify({'success': False, 'message': '目标项目不能在源项目列表中'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': '数据库连接失败'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 验证目标项目权限
        cursor.execute("SELECT created_by, name FROM projects WHERE id = %s", (target_project_id,))
        target_project = cursor.fetchone()
        
        if not target_project:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '目标项目不存在'}), 404
        
        if target_project['created_by'] != user_id:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '无权限操作目标项目'}), 403
        
        total_merged = 0
        merged_projects = []
        
        for source_project_id in source_project_ids:
            # 验证源项目权限
            cursor.execute("SELECT created_by, name FROM projects WHERE id = %s", (source_project_id,))
            source_project = cursor.fetchone()
            
            if not source_project:
                continue  # 跳过不存在的项目
            
            if source_project['created_by'] != user_id:
                continue  # 跳过无权限的项目
            
            # 获取源项目的所有资产
            cursor.execute("SELECT * FROM assets WHERE project_id = %s", (source_project_id,))
            assets = cursor.fetchall()
            
            # 复制资产到目标项目
            for asset in assets:
                # 检查目标项目中是否已存在相同的资产（根据 asset_value 判断）
                cursor.execute("""
                    SELECT id FROM assets 
                    WHERE project_id = %s AND asset_value = %s
                    LIMIT 1
                """, (target_project_id, asset['asset_value']))
                
                existing_asset = cursor.fetchone()
                
                if existing_asset:
                    # 如果已存在，更新资产信息（合并数据）
                    update_fields = []
                    update_values = []
                    
                    # 合并标签
                    if asset['tags']:
                        cursor.execute("SELECT tags FROM assets WHERE id = %s", (existing_asset['id'],))
                        existing_tags = cursor.fetchone()['tags'] or ''
                        merged_tags = list(set((existing_tags + ',' + asset['tags']).split(',')))
                        merged_tags = ','.join([t.strip() for t in merged_tags if t.strip()])
                        update_fields.append("tags = %s")
                        update_values.append(merged_tags)
                    
                    # 合并备注
                    if asset['notes']:
                        cursor.execute("SELECT notes FROM assets WHERE id = %s", (existing_asset['id'],))
                        existing_notes = cursor.fetchone()['notes'] or ''
                        merged_notes = existing_notes + '\n---\n' + asset['notes'] if existing_notes else asset['notes']
                        update_fields.append("notes = %s")
                        update_values.append(merged_notes)
                    
                    # 更新其他可能为空的字段
                    if asset['ip'] and not existing_asset.get('ip'):
                        update_fields.append("ip = %s")
                        update_values.append(asset['ip'])
                    
                    if asset['port'] and not existing_asset.get('port'):
                        update_fields.append("port = %s")
                        update_values.append(asset['port'])
                    
                    if asset['protocol'] and not existing_asset.get('protocol'):
                        update_fields.append("protocol = %s")
                        update_values.append(asset['protocol'])
                    
                    if asset['title'] and not existing_asset.get('title'):
                        update_fields.append("title = %s")
                        update_values.append(asset['title'])
                    
                    if asset['test_result'] and not existing_asset.get('test_result'):
                        update_fields.append("test_result = %s")
                        update_values.append(asset['test_result'])
                    
                    # 如果有需要更新的字段，执行更新
                    if update_fields:
                        update_values.append(existing_asset['id'])
                        cursor.execute(f"""
                            UPDATE assets 
                            SET {', '.join(update_fields)}
                            WHERE id = %s
                        """, tuple(update_values))
                else:
                    # 如果不存在，插入新资产
                    cursor.execute("""
                        INSERT INTO assets 
                        (project_id, asset_type, asset_value, ip, port, protocol, url, title, content_length, 
                         status, risk_level, tags, notes, test_result, created_by)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        target_project_id,
                        asset['asset_type'],
                        asset['asset_value'],
                        asset['ip'],
                        asset['port'],
                        asset['protocol'],
                        asset['url'],
                        asset['title'],
                        asset['content_length'],
                        asset['status'],
                        asset['risk_level'],
                        asset['tags'],
                        asset['notes'],
                        asset['test_result'],
                        user_id
                    ))
                
                total_merged += 1
            
            merged_projects.append({
                'id': source_project_id,
                'name': source_project['name'],
                'asset_count': len(assets)
            })
            
            # 如果需要删除源项目
            if delete_source:
                cursor.execute("DELETE FROM projects WHERE id = %s", (source_project_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'成功合并 {len(merged_projects)} 个项目，共 {total_merged} 个资产',
            'merged_projects': merged_projects,
            'target_project': {
                'id': target_project_id,
                'name': target_project['name']
            }
        })
    
    except Exception as e:
        return jsonify({'success': False, 'message': f'合并项目失败: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=15060, debug=True)
